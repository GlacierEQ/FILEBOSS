"""
Metadata Extractors

This module provides specialized extractors for different file types to extract
metadata and content for the file organizer.
"""

import os
import re
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Tuple
from datetime import datetime
import mimetypes
import email
from email import policy
import xml.etree.ElementTree as ET
import zipfile
import json

# Third-party imports (make sure to add to requirements.txt)
try:
    import PyPDF2
    import pdfplumber
    from pdfminer.high_level import extract_text
    from PIL import Image, ExifTags
    from PIL.Exif import TAGS
    import olefile
    from pptx import Presentation
    from docx import Document
    from openpyxl import load_workbook
    import pytesseract
    from pdf2image import convert_from_path
    import pytz
    from dateutil.parser import parse as parse_date
    import magic
    HAS_DEPENDENCIES = True
except ImportError:
    HAS_DEPENDENCIES = False
    logging.warning("Some metadata extraction features may be limited. Install with 'pip install -r requirements.txt'")

logger = logging.getLogger(__name__)

class MetadataExtractionError(Exception):
    """Exception raised for errors in metadata extraction."""
    pass

class BaseMetadataExtractor:
    """Base class for metadata extractors."""
    
    @classmethod
    def extract_metadata(cls, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from a file."""
        raise NotImplementedError("Subclasses must implement extract_metadata")
    
    @classmethod
    def extract_text_content(cls, file_path: Path) -> str:
        """Extract text content from a file."""
        return ""
    
    @classmethod
    def get_creation_date(cls, file_path: Path) -> Optional[datetime]:
        """Get creation date from file metadata."""
        try:
            stat = file_path.stat()
            return datetime.fromtimestamp(stat.st_ctime)
        except Exception as e:
            logger.warning(f"Could not get creation date for {file_path}: {e}")
            return None
    
    @classmethod
    def get_modification_date(cls, file_path: Path) -> Optional[datetime]:
        """Get modification date from file metadata."""
        try:
            stat = file_path.stat()
            return datetime.fromtimestamp(stat.st_mtime)
        except Exception as e:
            logger.warning(f"Could not get modification date for {file_path}: {e}")
            return None

class PDFMetadataExtractor(BaseMetadataExtractor):
    """Extract metadata from PDF files."""
    
    @classmethod
    def extract_metadata(cls, file_path: Path) -> Dict[str, Any]:
        metadata = {}
        
        try:
            # Extract basic PDF info
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                if pdf.is_encrypted:
                    try:
                        pdf.decrypt('')  # Try empty string as password
                    except:
                        logger.warning(f"Could not decrypt PDF: {file_path}")
                        return metadata
                
                # Get document info
                info = pdf.metadata or {}
                for key, value in info.items():
                    if hasattr(value, 'encode'):  # Skip non-string values
                        metadata[key.strip('/')] = value
                
                # Get page count
                metadata['page_count'] = len(pdf.pages)
                
            # Extract text for additional metadata
            try:
                with pdfplumber.open(file_path) as pdf:
                    # Try to get title from first page
                    if pdf.pages and not metadata.get('title'):
                        first_page = pdf.pages[0]
                        text = first_page.extract_text()
                        if text:
                            # Simple heuristic: first non-empty line could be the title
                            lines = [line.strip() for line in text.split('\n') if line.strip()]
                            if lines:
                                metadata['title'] = lines[0][:200]  # Limit title length
            except Exception as e:
                logger.warning(f"Could not extract text from PDF for metadata: {e}")
                
        except Exception as e:
            logger.error(f"Error extracting PDF metadata from {file_path}: {e}")
            
        return metadata
    
    @classmethod
    def extract_text_content(cls, file_path: Path) -> str:
        """Extract text content from PDF."""
        try:
            text = extract_text(str(file_path))
            return text.strip() if text else ""
        except Exception as e:
            logger.warning(f"Could not extract text from PDF {file_path}: {e}")
            return ""

class OfficeDocxMetadataExtractor(BaseMetadataExtractor):
    """Extract metadata from Word DOCX files."""
    
    @classmethod
    def extract_metadata(cls, file_path: Path) -> Dict[str, Any]:
        metadata = {}
        
        try:
            doc = Document(file_path)
            
            # Core properties
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.last_modified_by:
                metadata['last_modified_by'] = core_props.last_modified_by
            if core_props.created:
                metadata['created'] = core_props.created
            if core_props.modified:
                metadata['modified'] = core_props.modified
            
            # Basic stats
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['tables_count'] = len(doc.tables)
            
            # Try to get title from first paragraph if not in metadata
            if not metadata.get('title') and doc.paragraphs:
                first_para = doc.paragraphs[0].text.strip()
                if first_para:
                    metadata['title'] = first_para[:200]  # Limit title length
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata from {file_path}: {e}")
            
        return metadata
    
    @classmethod
    def extract_text_content(cls, file_path: Path) -> str:
        """Extract text content from DOCX."""
        try:
            doc = Document(file_path)
            return '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            logger.warning(f"Could not extract text from DOCX {file_path}: {e}")
            return ""

class ExcelMetadataExtractor(BaseMetadataExtractor):
    """Extract metadata from Excel files."""
    
    @classmethod
    def extract_metadata(cls, file_path: Path) -> Dict[str, Any]:
        metadata = {}
        
        try:
            wb = load_workbook(filename=file_path, read_only=True, data_only=True)
            
            # Basic info
            metadata['sheet_names'] = wb.sheetnames
            metadata['active_sheet'] = wb.active.title if wb.active else None
            
            # Sheet stats
            sheet_stats = {}
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_stats[sheet_name] = {
                    'rows': sheet.max_row,
                    'columns': sheet.max_column
                }
            metadata['sheet_stats'] = sheet_stats
            
            # Try to get title from first cell of first sheet
            if wb.sheetnames:
                first_sheet = wb[wb.sheetnames[0]]
                if first_sheet['A1'].value:
                    metadata['title'] = str(first_sheet['A1'].value)[:200]
            
        except Exception as e:
            logger.error(f"Error extracting Excel metadata from {file_path}: {e}")
            
        return metadata
    
    @classmethod
    def extract_text_content(cls, file_path: Path) -> str:
        """Extract text content from Excel."""
        try:
            wb = load_workbook(filename=file_path, read_only=True, data_only=True)
            content = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_content = [f"--- Sheet: {sheet_name} ---"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_content = []
                    for cell in row:
                        if cell is not None:
                            row_content.append(str(cell).strip())
                    if row_content:  # Only add non-empty rows
                        sheet_content.append(' | '.join(row_content))
                
                content.append('\n'.join(sheet_content))
            
            return '\n\n'.join(content)
            
        except Exception as e:
            logger.warning(f"Could not extract text from Excel {file_path}: {e}")
            return ""

class ImageMetadataExtractor(BaseMetadataExtractor):
    """Extract metadata from image files."""
    
    @classmethod
    def extract_metadata(cls, file_path: Path) -> Dict[str, Any]:
        metadata = {}
        
        try:
            with Image.open(file_path) as img:
                # Basic image info
                metadata['format'] = img.format
                metadata['mode'] = img.mode
                metadata['size'] = img.size  # (width, height)
                
                # Extract EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif_data = {}
                    for tag, value in img._getexif().items():
                        if tag in ExifTags.TAGS:
                            tag_name = ExifTags.TAGS[tag]
                            # Convert some common EXIF tags to more readable format
                            if tag_name == 'DateTimeOriginal':
                                try:
                                    exif_data['date_taken'] = datetime.strptime(
                                        value, '%Y:%m:%d %H:%M:%S')
                                except (ValueError, TypeError):
                                    exif_data[tag_name] = str(value)
                            elif tag_name in ('GPSInfo', 'GPSInfoVersion'):
                                # Handle GPS data specially if needed
                                pass
                            else:
                                exif_data[tag_name] = str(value)
                    
                    if exif_data:
                        metadata['exif'] = exif_data
                
                # Try to extract text using OCR if no other metadata
                if not metadata.get('title') and not metadata.get('description'):
                    try:
                        text = pytesseract.image_to_string(img)
                        if text.strip():
                            metadata['extracted_text'] = text.strip()
                    except Exception as e:
                        logger.debug(f"OCR not available or failed on {file_path}: {e}")
                
        except Exception as e:
            logger.error(f"Error extracting image metadata from {file_path}: {e}")
            
        return metadata
    
    @classmethod
    def extract_text_content(cls, file_path: Path) -> str:
        """Extract text content from image using OCR."""
        try:
            with Image.open(file_path) as img:
                text = pytesseract.image_to_string(img)
                return text.strip() if text else ""
        except Exception as e:
            logger.warning(f"Could not extract text from image {file_path}: {e}")
            return ""

class EmailMetadataExtractor(BaseMetadataExtractor):
    """Extract metadata from email files (EML, MSG)."""
    
    @classmethod
    def extract_metadata(cls, file_path: Path) -> Dict[str, Any]:
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                msg = email.message_from_file(f, policy=policy.default)
                
                # Basic headers
                for header in ['subject', 'from', 'to', 'cc', 'bcc', 'date', 
                             'message-id', 'in-reply-to', 'references']:
                    value = msg.get(header)
                    if value:
                        metadata[header] = str(value)
                
                # Parse date if available
                if 'date' in metadata:
                    try:
                        metadata['parsed_date'] = parse_date(metadata['date']).isoformat()
                    except (ValueError, AttributeError):
                        pass
                
                # Count attachments
                attachment_count = sum(1 for part in msg.walk() 
                                     if part.get_content_disposition() == 'attachment')
                if attachment_count > 0:
                    metadata['attachment_count'] = attachment_count
                
                # Get email body
                body = cls._get_email_body(msg)
                if body:
                    metadata['body_preview'] = body[:500]  # First 500 chars
        
        except Exception as e:
            logger.error(f"Error extracting email metadata from {file_path}: {e}")
            
        return metadata
    
    @classmethod
    def _get_email_body(cls, msg) -> str:
        """Extract the main email body text."""
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get('Content-Disposition'))
                
                # Skip attachments
                if 'attachment' in content_disposition:
                    continue
                    
                # Look for text/plain or text/html parts
                if content_type == 'text/plain':
                    try:
                        return part.get_payload(decode=True).decode('utf-8', 'ignore')
                    except:
                        continue
                
                # Fallback to HTML if no plain text
                if content_type == 'text/html':
                    try:
                        return part.get_payload(decode=True).decode('utf-8', 'ignore')
                    except:
                        continue
        else:
            # Not multipart, just return the payload
            try:
                return msg.get_payload(decode=True).decode('utf-8', 'ignore')
            except:
                pass
                
        return ""
    
    @classmethod
    def extract_text_content(cls, file_path: Path) -> str:
        """Extract text content from email."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                msg = email.message_from_file(f, policy=policy.default)
                return cls._get_email_body(msg)
        except Exception as e:
            logger.warning(f"Could not extract text from email {file_path}: {e}")
            return ""

class GenericMetadataExtractor(BaseMetadataExtractor):
    """Generic metadata extractor for unsupported file types."""
    
    @classmethod
    def extract_metadata(cls, file_path: Path) -> Dict[str, Any]:
        """Extract basic metadata for any file type."""
        metadata = {}
        
        try:
            # Basic file info
            stat = file_path.stat()
            metadata['size'] = stat.st_size
            metadata['created'] = datetime.fromtimestamp(stat.st_ctime).isoformat()
            metadata['modified'] = datetime.fromtimestamp(stat.st_mtime).isoformat()
            
            # Try to determine MIME type
            mime = magic.Magic(mime=True)
            metadata['mime_type'] = mime.from_file(str(file_path))
            
            # Try to extract text content for indexing
            try:
                text = cls.extract_text_content(file_path)
                if text:
                    metadata['extracted_text'] = text[:1000]  # First 1000 chars
            except Exception as e:
                logger.debug(f"Could not extract text from {file_path}: {e}")
                
        except Exception as e:
            logger.error(f"Error extracting generic metadata from {file_path}: {e}")
            
        return metadata
    
    @classmethod
    def extract_text_content(cls, file_path: Path) -> str:
        """Try to extract text from any file."""
        try:
            # Try to read as text file first
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
        except:
            return ""

# Registry of extractors by file extension
EXTRACTOR_REGISTRY = {
    # PDF
    '.pdf': PDFMetadataExtractor,
    
    # Microsoft Office
    '.docx': OfficeDocxMetadataExtractor,
    '.doc': OfficeDocxMetadataExtractor,  # Note: .doc requires additional libraries
    '.xlsx': ExcelMetadataExtractor,
    '.xls': ExcelMetadataExtractor,
    '.pptx': None,  # TODO: Add PowerPoint extractor
    '.ppt': None,   # TODO: Add PowerPoint extractor
    
    # Images
    '.jpg': ImageMetadataExtractor,
    '.jpeg': ImageMetadataExtractor,
    '.png': ImageMetadataExtractor,
    '.gif': ImageMetadataExtractor,
    '.bmp': ImageMetadataExtractor,
    '.tiff': ImageMetadataExtractor,
    '.webp': ImageMetadataExtractor,
    
    # Email
    '.eml': EmailMetadataExtractor,
    '.msg': EmailMetadataExtractor,  # Note: .msg requires additional parsing
    
    # Archives
    '.zip': None,  # TODO: Add archive extractor
    '.rar': None,
    '.7z': None,
    '.tar': None,
    '.gz': None,
    '.bz2': None,
}

def get_metadata_extractor(file_path: Path) -> BaseMetadataExtractor:
    """Get the appropriate metadata extractor for a file."""
    ext = file_path.suffix.lower()
    return EXTRACTOR_REGISTRY.get(ext, GenericMetadataExtractor)

def extract_metadata(file_path: Path) -> Dict[str, Any]:
    """Extract metadata from a file using the appropriate extractor."""
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extractor = get_metadata_extractor(file_path)
    return extractor.extract_metadata(file_path)

def extract_text_content(file_path: Path) -> str:
    """Extract text content from a file using the appropriate extractor."""
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    extractor = get_metadata_extractor(file_path)
    return extractor.extract_text_content(file_path)
