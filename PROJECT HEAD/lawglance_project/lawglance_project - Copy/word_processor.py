from docx import Document

class WordProcessor:
    """Utility class for word processing tasks."""

    @staticmethod
    def read_document(file_path):
        """Read a Word document and return its text content."""
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])

    @staticmethod
    def write_document(file_path, content):
        """Write content to a Word document."""
        doc = Document()
        doc.add_paragraph(content)
        doc.save(file_path)
