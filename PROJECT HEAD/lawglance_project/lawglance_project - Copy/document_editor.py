import os
import re
import difflib
import docx
from docx.shared import Pt, RGBColor
import docx.enum.text as text
from docx.enum.style import WD_STYLE_TYPE
import logging

try:
    import win32com.client
    import pythoncom
    COM_AVAILABLE = True
except ImportError:
    COM_AVAILABLE = False

class DocumentEditor:
    """Edits documents based on natural language instructions."""
    
    def __init__(self):
        """Initialize the document editor."""
        self.logger = logging.getLogger("lawglance.document_editor")
        
        # Available edit operations
        self.edit_operations = {
            "replace": self._replace_text,
            "insert": self._insert_text,
            "delete": self._delete_text,
            "format": self._format_text,
            "add_section": self._add_section,
            "move": self._move_text
        }
        
        # Formatting styles dictionary
        self.formatting = {
            "bold": lambda p, r: setattr(r.font, "bold", True),
            "italic": lambda p, r: setattr(r.font, "italic", True),
            "underline": lambda p, r: setattr(r.font, "underline", True),
            "red": lambda p, r: setattr(r.font.color, "rgb", RGBColor(255, 0, 0)),
            "blue": lambda p, r: setattr(r.font.color, "rgb", RGBColor(0, 0, 255)),
            "green": lambda p, r: setattr(r.font.color, "rgb", RGBColor(0, 128, 0)),
            "large": lambda p, r: setattr(r.font, "size", Pt(14)),
            "small": lambda p, r: setattr(r.font, "size", Pt(8)),
            "heading": lambda p, r: self._apply_heading_style(p),
            "normal": lambda p, r: self._apply_normal_style(p),
        }

    def edit_document(self, file_path, instructions):
        """Edit a document based on natural language instructions.
        
        Args:
            file_path: Path to the document
            instructions: Natural language description of edits to perform
            
        Returns:
            Status message indicating edits performed
        """
        if not os.path.exists(file_path):
            return f"Error: File {file_path} not found"
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self._edit_docx(file_path, instructions)
        elif file_ext == '.doc' and COM_AVAILABLE:
            return self._edit_doc(file_path, instructions)
        elif file_ext in ['.txt', '.md']:
            return self._edit_text_file(file_path, instructions)
        else:
            return f"Error: Unsupported file format {file_ext}"
    
    def _edit_docx(self, file_path, instructions):
        """Edit a .docx file."""
        try:
            doc = docx.Document(file_path)
            
            # Parse instruction to determine operation
            operation, params = self._parse_instructions(instructions)
            
            if operation not in self.edit_operations:
                return f"Error: Unsupported operation '{operation}'"
                
            # Perform the operation
            result = self.edit_operations[operation](doc, **params)
            
            # Save with a backup
            backup_path = f"{file_path}.backup"
            if not os.path.exists(backup_path):
                import shutil
                shutil.copy2(file_path, backup_path)
                
            doc.save(file_path)
            return result
            
        except Exception as e:
            self.logger.error(f"Error editing document: {str(e)}")
            return f"Error editing document: {str(e)}"
    
    def _edit_doc(self, file_path, instructions):
        """Edit a .doc file using COM interface."""
        try:
            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Convert to docx temporarily
            temp_docx = f"{file_path}x"
            doc = word.Documents.Open(file_path)
            doc.SaveAs(temp_docx, 16)  # 16 = wdFormatDocumentDefault (docx)
            doc.Close()
            
            # Edit the docx version
            result = self._edit_docx(temp_docx, instructions)
            
            # Convert back to original format and cleanup
            doc = word.Documents.Open(temp_docx)
            doc.SaveAs(file_path, 0)  # 0 = wdFormatDocument (doc)
            doc.Close()
            
            word.Quit()
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
                
            return result
            
        except Exception as e:
            self.logger.error(f"Error editing .doc file: {str(e)}")
            return f"Error editing .doc file: {str(e)}"
        finally:
            pythoncom.CoUninitialize()
    
    def _edit_text_file(self, file_path, instructions):
        """Edit a simple text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
                
            operation, params = self._parse_instructions(instructions)
            
            # Handle operations for text files
            if operation == "replace":
                old_text = params.get("old_text", "")
                new_text = params.get("new_text", "")
                
                if old_text and new_text:
                    content = content.replace(old_text, new_text)
                    
            elif operation == "insert":
                position_text = params.get("position", "")
                new_text = params.get("text", "")
                
                if position_text and new_text:
                    if position_text.lower() == "start":
                        content = new_text + content
                    elif position_text.lower() == "end":
                        content = content + new_text
                    else:
                        # Insert after the position text
                        content = content.replace(position_text, position_text + new_text)
                        
            elif operation == "delete":
                text_to_delete = params.get("text", "")
                if text_to_delete:
                    content = content.replace(text_to_delete, "")
                    
            # Create backup
            backup_path = f"{file_path}.backup"
            if not os.path.exists(backup_path):
                with open(backup_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                    
            # Write updated content
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
                
            return f"Successfully edited {file_path}"
            
        except Exception as e:
            self.logger.error(f"Error editing text file: {str(e)}")
            return f"Error editing text file: {str(e)}"
    
    def _parse_instructions(self, instructions):
        """Parse natural language instructions into operation and parameters."""
        instructions = instructions.lower()
        
        # Try to identify the operation
        operation = None
        params = {}
        
        # Check for replace operation
        replace_patterns = [
            r"replace\s+['\"](.+?)['\"]\s+with\s+['\"](.+?)['\"]",
            r"change\s+['\"](.+?)['\"]\s+to\s+['\"](.+?)['\"]"
        ]
        
        for pattern in replace_patterns:
            match = re.search(pattern, instructions)
            if match:
                operation = "replace"
                params["old_text"] = match.group(1)
                params["new_text"] = match.group(2)
                break
                
        # Check for insert operation
        if not operation:
            insert_patterns = [
                r"insert\s+['\"](.+?)['\"]\s+(?:at|after|before)\s+['\"](.+?)['\"]",
                r"add\s+['\"](.+?)['\"]\s+(?:at|after|before)\s+['\"](.+?)['\"]"
            ]
            
            for pattern in insert_patterns:
                match = re.search(pattern, instructions)
                if match:
                    operation = "insert"
                    params["text"] = match.group(1)
                    params["position"] = match.group(2)
                    break
                    
            # Check for insert at beginning/end
            if not operation:
                start_pattern = r"(?:insert|add)\s+['\"](.+?)['\"]\s+at\s+(?:the\s+)?(?:beginning|start)"
                end_pattern = r"(?:insert|add)\s+['\"](.+?)['\"]\s+at\s+(?:the\s+)?end"
                
                start_match = re.search(start_pattern, instructions)
                end_match = re.search(end_pattern, instructions)
                
                if start_match:
                    operation = "insert"
                    params["text"] = start_match.group(1)
                    params["position"] = "start"
                elif end_match:
                    operation = "insert"
                    params["text"] = end_match.group(1)
                    params["position"] = "end"
        
        # Check for delete operation
        if not operation:
            delete_patterns = [
                r"(?:delete|remove)\s+['\"](.+?)['\"]",
                r"(?:erase|eliminate)\s+['\"](.+?)['\"]"
            ]
            
            for pattern in delete_patterns:
                match = re.search(pattern, instructions)
                if match:
                    operation = "delete"
                    params["text"] = match.group(1)
                    break
        
        # Check for format operation
        if not operation:
            format_patterns = [
                r"(?:format|make|set)\s+['\"](.+?)['\"]\s+(?:as|to)\s+(\w+)",
                r"(\w+)\s+['\"](.+?)['\"]"  # Bold "this text"
            ]
            
            for pattern in format_patterns:
                match = re.search(pattern, instructions)
                if match and match.group(2).lower() in self.formatting:
                    operation = "format"
                    params["text"] = match.group(1)
                    params["style"] = match.group(2).lower()
                    break
                elif match and match.group(1).lower() in self.formatting:
                    operation = "format"
                    params["text"] = match.group(2)
                    params["style"] = match.group(1).lower()
                    break
        
        # Check for add section operation
        if not operation:
            section_patterns = [
                r"add\s+(?:a\s+)?(?:new\s+)?section\s+['\"](.+?)['\"]\s+with\s+content\s+['\"](.+?)['\"]",
                r"create\s+(?:a\s+)?(?:new\s+)?section\s+['\"](.+?)['\"]\s+with\s+['\"](.+?)['\"]"
            ]
            
            for pattern in section_patterns:
                match = re.search(pattern, instructions)
                if match:
                    operation = "add_section"
                    params["title"] = match.group(1)
                    params["content"] = match.group(2)
                    break
        
        # Check for move operation
        if not operation:
            move_patterns = [
                r"move\s+['\"](.+?)['\"]\s+(?:before|after)\s+['\"](.+?)['\"]",
                r"relocate\s+['\"](.+?)['\"]\s+(?:before|after)\s+['\"](.+?)['\"]"
            ]
            
            for pattern in move_patterns:
                match = re.search(pattern, instructions)
                if match:
                    operation = "move"
                    params["text"] = match.group(1)
                    params["target"] = match.group(2)
                    params["position"] = "after" if "after" in instructions else "before"
                    break
        
        # Default to replace if we couldn't determine operation
        if not operation:
            return "replace", {"old_text": "", "new_text": ""}
            
        return operation, params
    
    def _replace_text(self, doc, old_text="", new_text=""):
        """Replace text in a document."""        
        if not old_text or not new_text:
            return "Error: Both old and new text must be provided for replacement"
            
        replacement_count = 0
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replacement_count += 1
        
        # Replace in tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    replacement_count += 1
        
        if replacement_count > 0:
            return f"Replaced '{old_text}' with '{new_text}' {replacement_count} times"
        else:
            return f"Text '{old_text}' not found in document"
    
    def _insert_text(self, doc, text="", position="end"):
        """Insert text at specified position."""        
        if not text:
            return "Error: Text to insert must be provided"
            
        if position == "start":
            # Insert at beginning of document
            para = doc.paragraphs[0]
            para.insert_paragraph_before(text)
            return f"Inserted text at the beginning of the document"
            
        elif position == "end":
            # Insert at end of document
            doc.add_paragraph(text)
            return f"Inserted text at the end of the document"
            
        else:
            # Insert after specified text
            for i, para in enumerate(doc.paragraphs):
                if position in para.text:
                    # If exact match to paragraph text, insert after this paragraph
                    if position == para.text.strip():
                        doc.paragraphs[i].insert_paragraph_before(text)
                        return f"Inserted text after '{position}'"
                    
                    # Try to insert within paragraph
                    for run in para.runs:
                        if position in run.text:
                            idx = run.text.find(position) + len(position)
                            before = run.text[:idx]
                            after = run.text[idx:]
                            run.text = before
                            
                            new_run = para.add_run(text + after)
                            # Copy formatting from original run
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.size = run.font.size
                            
                            return f"Inserted text after '{position}'"
            
            return f"Position '{position}' not found in document"
    
    def _delete_text(self, doc, text=""):
        """Delete text from document."""        
        if not text:
            return "Error: Text to delete must be provided"
            
        deletion_count = 0
        
        # Delete from paragraphs
        for para in doc.paragraphs:
            if text in para.text:
                for run in para.runs:
                    if text in run.text:
                        run.text = run.text.replace(text, "")
                        deletion_count += 1
        
        # Delete from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if text in para.text:
                            for run in para.runs:
                                if text in run.text:
                                    run.text = run.text.replace(text, "")
                                    deletion_count += 1
        
        if deletion_count > 0:
            return f"Deleted '{text}' {deletion_count} times"
        else:
            return f"Text '{text}' not found in document"
    
    def _format_text(self, doc, text="", style=""):
        """Apply formatting to text."""        
        if not text or not style:
            return "Error: Both text and style must be provided for formatting"
            
        if style not in self.formatting:
            return f"Error: Unsupported formatting style '{style}'"
            
        format_count = 0
        formatter = self.formatting[style]
        
        # Format paragraphs
        for para in doc.paragraphs:
            if text in para.text:
                if text == para.text.strip():
                    # Apply to entire paragraph
                    formatter(para, para.runs[0] if para.runs else None)
                    format_count += 1
                else:
                    # Apply to specific runs containing the text
                    for run in para.runs:
                        if text in run.text:
                            formatter(para, run)
                            format_count += 1
        
        if format_count > 0:
            return f"Applied {style} formatting to '{text}' {format_count} times"
        else:
            return f"Text '{text}' not found in document"
    
    def _add_section(self, doc, title="", content=""):
        """Add a new section with title and content."""        
        if not title:
            return "Error: Section title must be provided"
            
        # Add title
        heading = doc.add_heading(title, level=2)
        
        # Add content if provided
        if content:
            doc.add_paragraph(content)
            
        return f"Added new section '{title}'"
    
    def _move_text(self, doc, text="", target="", position="after"):
        """Move text to before or after the target text."""        
        if not text or not target:
            return "Error: Both source text and target position must be provided"
            
        # First find and store the text to move
        found_text = False
        text_runs = []
        
        # Find the text to move and its formatting
        for para in doc.paragraphs:
            if text in para.text:
                for run in para.runs:
                    if text in run.text:
                        # Save the text and its formatting
                        text_runs.append({
                            "text": text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                            "size": run.font.size
                        })
                        found_text = True
                        break
                if found_text:
                    break
        
        if not found_text:
            return f"Text '{text}' not found in document"
            
        # Delete the original occurrence
        self._delete_text(doc, text)
        
        # Now insert at the target position
        if position == "after":
            for para in doc.paragraphs:
                if target in para.text:
                    idx = para.text.find(target) + len(target)
                    before = para.text[:idx]
                    after = para.text[idx:]
                    
                    # Clear the paragraph
                    for run in para.runs:
                        run.clear()
                    
                    # Add back the content with the moved text
                    para.add_run(before)
                    moved_run = para.add_run(text)
                    
                    # Apply stored formatting
                    if text_runs:
                        moved_run.bold = text_runs[0]["bold"]
                        moved_run.italic = text_runs[0]["italic"]
                        moved_run.underline = text_runs[0]["underline"]
                        moved_run.font.size = text_runs[0]["size"]
                        
                    para.add_run(after)
                    return f"Moved '{text}' after '{target}'"
        else:  # position == "before"
            for para in doc.paragraphs:
                if target in para.text:
                    idx = para.text.find(target)
                    before = para.text[:idx]
                    after = para.text[idx:]
                    
                    # Clear the paragraph
                    for run in para.runs:
                        run.clear()
                    
                    # Add back the content with the moved text
                    para.add_run(before)
                    moved_run = para.add_run(text)
                    
                    # Apply stored formatting
                    if text_runs:
                        moved_run.bold = text_runs[0]["bold"]
                        moved_run.italic = text_runs[0]["italic"]
                        moved_run.underline = text_runs[0]["underline"]
                        moved_run.font.size = text_runs[0]["size"]
                        
                    para.add_run(after)
                    return f"Moved '{text}' before '{target}'"
        
        return f"Target '{target}' not found in document"
    
    def _apply_heading_style(self, paragraph):
        """Apply heading style to a paragraph."""        
        paragraph.style = "Heading 2"
        
    def _apply_normal_style(self, paragraph):
        """Apply normal style to a paragraph."""        
        paragraph.style = "Normal"
