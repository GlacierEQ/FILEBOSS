"""
Fixed methods for the document editor with enhanced error handling.
This module provides improved implementations that address edge cases.
"""
import os
import re
import logging
from typing import Dict, Any, Optional, Tuple, List, Union
from enum import Enum
from dataclasses import dataclass

import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import docx.enum.text as text_enums
from docx.enum.style import WD_STYLE_TYPE
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# Set up logging
logger = logging.getLogger("lawglance.editor_fixes")

class ParsingError(Exception):
    """Exception raised when parsing instructions fails."""
    pass

class TextNotFoundError(Exception):
    """Exception raised when text is not found in document."""
    pass

class ElementNotFoundError(Exception):
    """Exception raised when a paragraph or run is not found."""
    pass

class FormatNotSupportedError(Exception):
    """Exception raised when a formatting option is not supported."""
    pass

@dataclass
class EditResult:
    """Result of a document editing operation."""
    success: bool = True
    message: str = ""
    count: int = 0
    error: Optional[str] = None
    details: Optional[Dict[str, Any]] = None
    
    def __str__(self) -> str:
        """String representation of the edit result."""
        if not self.success:
            return f"Error: {self.error}"
        return self.message

class EditorFixes:
    """Fixed methods for document editor."""
    
    def __init__(self):
        """Initialize the editor fixes."""
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def parse_instructions(self, instructions: str) -> Tuple[str, Dict[str, Any]]:
        """
        Fixed version of _parse_instructions that provides detailed error messages.
        
        Args:
            instructions: Natural language instructions
            
        Returns:
            A tuple of (operation, parameters)
            
        Raises:
            ParsingError: If instructions cannot be parsed
        """
        if not instructions:
            raise ParsingError("Instructions cannot be empty")
            
        instructions = instructions.lower()
        
        # Try to identify the operation and parameters
        try:
            # Parse for replace operation
            replace_match = re.search(r"replace\s+['\"](.+?)['\"]\s+with\s+['\"](.+?)['\"]", instructions)
            if replace_match:
                return "replace", {
                    "old_text": replace_match.group(1),
                    "new_text": replace_match.group(2)
                }
                
            # Parse for insert operation
            insert_match = re.search(r"insert\s+['\"](.+?)['\"]\s+(?:at|after|before)\s+['\"](.+?)['\"]", instructions)
            if insert_match:
                return "insert", {
                    "text": insert_match.group(1),
                    "position": insert_match.group(2)
                }
                
            # Parse for start/end insertion
            start_match = re.search(r"insert\s+['\"](.+?)['\"](?:\s+at)?\s+(?:the\s+)?beginning", instructions)
            if start_match:
                return "insert", {
                    "text": start_match.group(1),
                    "position": "start"
                }
                
            end_match = re.search(r"insert\s+['\"](.+?)['\"](?:\s+at)?\s+(?:the\s+)?end", instructions)
            if end_match:
                return "insert", {
                    "text": end_match.group(1),
                    "position": "end"
                }
                
            # Parse for delete operation
            delete_match = re.search(r"(?:delete|remove)\s+['\"](.+?)['\"]", instructions)
            if delete_match:
                return "delete", {
                    "text": delete_match.group(1)
                }
            
            # Parse for format operation
            format_match = re.search(r"(?:format|make|set)\s+['\"](.+?)['\"]\s+(?:as|to)\s+(\w+)", instructions)
            if format_match:
                return "format", {
                    "text": format_match.group(1),
                    "style": format_match.group(2)
                }
                
            # Failed to parse instructions
            self.logger.error(f"Failed to parse instructions: '{instructions}'")
            raise ParsingError(
                f"Could not understand instructions: '{instructions}'. "
                f"Please use a supported format like 'replace \"text\" with \"new text\"' or "
                f"'insert \"text\" at \"position\"'."
            )
            
        except Exception as e:
            self.logger.error(f"Error parsing instructions '{instructions}': {str(e)}")
            raise ParsingError(
                f"Error parsing instructions: {str(e)}. Please check the format of your instructions."
            )
    
    def format_text(self, doc: Document, text: str, style: str) -> EditResult:
        """
        Fixed version of _format_text that handles empty text runs.
        
        Args:
            doc: The document object to edit
            text: Text to format
            style: Formatting style to apply
            
        Returns:
            EditResult with operation status
        """
        if not text:
            return EditResult(
                success=False,
                error="Text to format must be provided"
            )
            
        if not style:
            return EditResult(
                success=False,
                error="Formatting style must be specified"
            )
            
        # Dictionary mapping style names to formatting functions
        format_handlers = {
            "bold": self._apply_bold,
            "italic": self._apply_italic,
            "underline": self._apply_underline,
            "heading1": lambda p: self._apply_heading_style(p, 1),
            "heading2": lambda p: self._apply_heading_style(p, 2),
            "heading3": lambda p: self._apply_heading_style(p, 3),
        }
        
        if style not in format_handlers:
            return EditResult(
                success=False,
                error=f"Unsupported formatting style: '{style}'. Supported styles are: {', '.join(format_handlers.keys())}"
            )
            
        format_count = 0
        format_function = format_handlers[style]
        
        # Track whether we found the text
        text_found = False
        
        # Format paragraphs
        for para in doc.paragraphs:
            if text in para.text:
                text_found = True
                if text == para.text.strip():
                    # Apply to entire paragraph
                    format_function(para)
                    format_count += 1
                else:
                    # Apply to specific runs containing the text
                    for run in para.runs:
                        # Skip empty runs
                        if not run.text:
                            continue
                            
                        if text in run.text:
                            format_function(run)
                            format_count += 1
        
        # Format tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if text in para.text:
                            text_found = True
                            if text == para.text.strip():
                                format_function(para)
                                format_count += 1
                            else:
                                # Apply to specific runs containing the text
                                for run in para.runs:
                                    # Skip empty runs
                                    if not run.text:
                                        continue
                                        
                                    if text in run.text:
                                        format_function(run)
                                        format_count += 1
        
        if not text_found:
            return EditResult(
                success=False,
                error=f"Text '{text}' not found in document",
                details={"text_to_format": text, "style": style}
            )
            
        if format_count == 0:
            # This can happen if we found the text but couldn't apply formatting
            # (e.g., due to empty runs or incompatible elements)
            return EditResult(
                success=False,
                error=f"Found '{text}' but could not apply formatting",
                details={"text_to_format": text, "style": style}
            )
            
        return EditResult(
            success=True,
            message=f"Applied {style} formatting to '{text}' {format_count} times",
            count=format_count
        )
    
    def move_text(self, doc: Document, text: str, target: str, position: str = "after") -> EditResult:
        """
        Fixed version of _move_text that handles case where text is not found.
        
        Args:
            doc: The document object to edit
            text: Text to move
            target: Target position text
            position: "before" or "after" target
            
        Returns:
            EditResult with operation status
        """
        if not text or not target:
            return EditResult(
                success=False,
                error="Both source text and target position must be provided"
            )
            
        if position not in ["before", "after"]:
            return EditResult(
                success=False,
                error="Position must be 'before' or 'after'"
            )
        
        # First find and store the text to move
        found_text = False
        text_runs = []
        
        # Phase 1: Find the text to move
        # Look through paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if text in para.text:
                for run_idx, run in enumerate(para.runs):
                    if text in run.text:
                        # Save the text and its formatting
                        text_runs.append({
                            "text": text,
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                            "size": run.font.size,
                            "color": run.font.color.rgb if run.font.color.rgb else None,
                            "font": run.font.name,
                            "location": ("para", para_idx, run_idx)
                        })
                        found_text = True
                        break
                if found_text:
                    break
        
        # Look through tables if not found in paragraphs
        if not found_text:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            if text in para.text:
                                for run_idx, run in enumerate(para.runs):
                                    if text in run.text:
                                        # Save the text and its formatting
                                        text_runs.append({
                                            "text": text,
                                            "bold": run.bold,
                                            "italic": run.italic,
                                            "underline": run.underline,
                                            "size": run.font.size,
                                            "color": run.font.color.rgb if run.font.color.rgb else None,
                                            "font": run.font.name,
                                            "location": ("table", table_idx, row_idx, cell_idx, para_idx, run_idx)
                                        })
                                        found_text = True
                                        break
                                if found_text:
                                    break
                            if found_text:
                                break
                        if found_text:
                            break
                    if found_text:
                        break
                if found_text:
                    break
        
        if not found_text:
            return EditResult(
                success=False,
                error=f"Text '{text}' not found in document",
                details={"text_to_move": text, "target": target, "position": position}
            )
        
        # Phase 2: Find the target text
        found_target = False
        target_location = None
        
        # Look for target in paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if target in para.text:
                target_location = ("para", para_idx)
                found_target = True
                break
        
        # Look for target in tables if not found in paragraphs
        if not found_target:
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            if target in para.text:
                                target_location = ("table", table_idx, row_idx, cell_idx, para_idx)
                                found_target = True
                                break
                        if found_target:
                            break
                    if found_target:
                        break
                if found_target:
                    break
        
        if not found_target:
            return EditResult(
                success=False,
                error=f"Target text '{target}' not found in document",
                details={"text_to_move": text, "target": target, "position": position}
            )
        
        # Phase 3: Delete the original text
        delete_result = self.delete_text(doc, text)
        if not delete_result.success:
            return delete_result
        
        # Phase 4: Insert the text at the target position
        if target_location[0] == "para":
            para = doc.paragraphs[target_location[1]]
            
            # If the target is the whole paragraph, insert before/after it
            if target == para.text.strip():
                if position == "after":
                    # Insert a new paragraph after it
                    new_para = doc.add_paragraph("")
                    # Move the new paragraph to the correct position
                    doc._body._body.insert(target_location[1] + 1, new_para._p)
                    for text_run in text_runs:
                        new_run = new_para.add_run(text_run["text"])
                        self._apply_run_formatting(new_run, text_run)
                else:  # position == "before"
                    # Insert a new paragraph before it
                    new_para = para.insert_paragraph_before("")
                    for text_run in text_runs:
                        new_run = new_para.add_run(text_run["text"])
                        self._apply_run_formatting(new_run, text_run)
            else:
                # Find the position within the paragraph
                target_idx = para.text.find(target)
                if position == "after":
                    target_idx += len(target)
                
                # Insert at the specified position
                for i, run in enumerate(para.runs):
                    if target in run.text:
                        idx = run.text.find(target)
                        if position == "after":
                            idx += len(target)
                            
                        # Split the run
                        before = run.text[:idx]
                        after = run.text[idx:]
                        run.text = before
                        
                        # Insert the moved text
                        for text_run in text_runs:
                            moved_run = para.add_run(text_run["text"])
                            self._apply_run_formatting(moved_run, text_run)
                        
                        # Add the remainder
                        if after:
                            remainder_run = para.add_run(after)
                            # Copy formatting from original run
                            remainder_run.bold = run.bold
                            remainder_run.italic = run.italic
                            remainder_run.underline = run.underline
                            remainder_run.font.size = run.font.size
                        break
        else:
            # Not implementing table cell editing here for brevity
            self.logger.warning("Moving text to a table cell is not implemented")
            return EditResult(
                success=False,
                error="Moving text to a table cell is not supported yet",
                details={"text_to_move": text, "target": target, "position": position}
            )
        
        return EditResult(
            success=True,
            message=f"Moved '{text}' {position} '{target}'",
            count=1
        )
    
    def delete_text(self, doc: Document, text: str) -> EditResult:
        """
        Fixed version of _delete_text with better error handling.
        
        Args:
            doc: The document object to edit
            text: Text to delete
            
        Returns:
            EditResult with operation status
        """
        if not text:
            return EditResult(
                success=False,
                error="Text to delete must be provided"
            )
            
        deletion_count = 0
        
        # Delete from paragraphs
        for para in doc.paragraphs:
            if text in para.text:
                for run in para.runs:
                    if text in run.text:
                        run.text = run.text.replace(text, "")
                        deletion_count += 1
        
        # Delete from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if text in para.text:
                            for run in para.runs:
                                if text in run.text:
                                    run.text = run.text.replace(text, "")
                                    deletion_count += 1
        
        if deletion_count == 0:
            return EditResult(
                success=False,
                error=f"Text '{text}' not found in document",
                details={"text_to_delete": text}
            )
        
        return EditResult(
            success=True,
            message=f"Deleted '{text}' {deletion_count} times",
            count=deletion_count
        )
    
    def insert_text(self, doc: Document, text: str, position: str) -> EditResult:
        """
        Fixed version of _insert_text that handles case where insertion point is not found.
        
        Args:
            doc: The document object to edit
            text: Text to insert
            position: Location to insert text ("start", "end", or specific text)
            
        Returns:
            EditResult with operation status
        """
        if not text:
            return EditResult(
                success=False,
                error="Text to insert must be provided"
            )
            
        if position == "start":
            # Insert at beginning of document
            if not doc.paragraphs:
                doc.add_paragraph(text)
            else:
                doc.paragraphs[0].insert_paragraph_before(text)
            return EditResult(
                success=True,
                message=f"Inserted text at the beginning of the document",
                count=1
            )
            
        elif position == "end":
            # Insert at end of document
            doc.add_paragraph(text)
            return EditResult(
                success=True,
                message=f"Inserted text at the end of the document",
                count=1
            )
            
        else:
            # Insert after specific text
            position_found = False
            
            # Check paragraphs
            for i, para in enumerate(doc.paragraphs):
                if position in para.text:
                    position_found = True
                    
                    # If exact match to paragraph text, insert after this paragraph
                    if position == para.text.strip():
                        # Create a new paragraph after the current one
                        new_para = doc.add_paragraph("")
                        # Move it to the correct position
                        doc._body._body.insert(i + 1, new_para._p)
                        new_para.add_run(text)
                        return EditResult(
                            success=True,
                            message=f"Inserted text after '{position}'",
                            count=1
                        )
                    
                    # Insert within paragraph
                    for run in para.runs:
                        if position in run.text:
                            idx = run.text.find(position) + len(position)
                            before = run.text[:idx]
                            after = run.text[idx:]
                            run.text = before
                            
                            # Add new text
                            new_run = para.add_run(text + after)
                            # Copy formatting from original run
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.size = run.font.size
                            
                            return EditResult(
                                success=True,
                                message=f"Inserted text after '{position}'",
                                count=1
                            )
            
            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if position in para.text:
                                position_found = True
                                
                                for run in para.runs:
                                    if position in run.text:
                                        idx = run.text.find(position) + len(position)
                                        before = run.text[:idx]
                                        after = run.text[idx:]
                                        run.text = before
                                        
                                        # Add new text
                                        new_run = para.add_run(text + after)
                                        # Copy formatting from original run
                                        new_run.bold = run.bold
                                        new_run.italic = run.italic
                                        new_run.underline = run.underline
                                        new_run.font.size = run.font.size
                                        
                                        return EditResult(
                                            success=True,
                                            message=f"Inserted text after '{position}' in table",
                                            count=1
                                        )
            
            if not position_found:
                return EditResult(
                    success=False,
                    error=f"Position text '{position}' not found in document",
                    details={"text_to_insert": text, "position": position}
                )
            
            # If we get here, we found the position text but couldn't insert after it
            return EditResult(
                success=False,
                error=f"Found position '{position}' but could not insert text after it",
                details={"text_to_insert": text, "position": position}
            )
    
    def apply_heading_style(self, doc: Document, text: str, level: int) -> EditResult:
        """
        Fixed version of _apply_heading_style that handles case where paragraph is not found.
        
        Args:
            doc: The document object to edit
            text: Text to format
            level: Heading level (1-9)
            
        Returns:
            EditResult with operation status
        """
        if not text:
            return EditResult(
                success=False,
                error="Text to format must be provided"
            )
            
        if not isinstance(level, int) or level < 1 or level > 9:
            return EditResult(
                success=False,
                error=f"Heading level must be an integer between 1 and 9, got {level}",
                details={"text": text, "level": level}
            )
            
        format_count = 0
        paragraph_found = False
        
        # Format paragraphs
        for para in doc.paragraphs:
            if text in para.text and text == para.text.strip():
                paragraph_found = True
                try:
                    para.style = f'Heading {level}'
                    format_count += 1
                except Exception as e:
                    self.logger.error(f"Error applying heading style: {e}")
                    return EditResult(
                        success=False,
                        error=f"Error applying heading style: {str(e)}",
                        details={"text": text, "level": level, "error": str(e)}
                    )
        
        # Format tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if text in para.text and text == para.text.strip():
                            paragraph_found = True
                            try:
                                para.style = f'Heading {level}'
                                format_count += 1
                            except Exception as e:
                                self.logger.error(f"Error applying heading style in table: {e}")
                                return EditResult(
                                    success=False,
                                    error=f"Error applying heading style in table: {str(e)}",
                                    details={"text": text, "level": level, "error": str(e)}
                                )
        
        if not paragraph_found:
            return EditResult(
                success=False,
                error=f"Text '{text}' not found as a whole paragraph in document",
                details={"text": text, "level": level}
            )
        
        if format_count == 0:
            # This can happen if the text is found but we couldn't apply the style
            return EditResult(
                success=False,
                error=f"Found '{text}' but could not apply heading style",
                details={"text": text, "level": level}
            )
        
        return EditResult(
            success=True,
            message=f"Applied Heading {level} style to '{text}' {format_count} times",
            count=format_count
        )
    
    def apply_font_style(self, doc: Document, text: str, font_name: str, font_size: Optional[int] = None) -> EditResult:
        """
        Fixed version of _apply_font_style that handles case where text run is not found.
        
        Args:
            doc: The document object to edit
            text: Text to format
            font_name: Font name
            font_size: Font size in points (optional)
            
        Returns:
            EditResult with operation status
        """
        if not text:
            return EditResult(
                success=False,
                error="Text to format must be provided"
            )
            
        if not font_name:
            return EditResult(
                success=False,
                error="Font name must be provided"
            )
            
        format_count = 0
        text_found = False
        
        # Format paragraphs
        for para in doc.paragraphs:
            if text in para.text:
                text_found = True
                
                for run in para.runs:
                    if text in run.text:
                        try:
                            run.font.name = font_name
                            if font_size is not None:
                                run.font.size = Pt(font_size)
                            format_count += 1
                        except Exception as e:
                            self.logger.error(f"Error applying font style: {e}")
                            return EditResult(
                                success=False,
                                error=f"Error applying font style: {str(e)}",
                                details={"text": text, "font_name": font_name, "font_size": font_size, "error": str(e)}
                            )
        
        # Format tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if text in para.text:
                            text_found = True
                            
                            for run in para.runs:
                                if text in run.text:
                                    try:
                                        run.font.name = font_name
                                        if font_size is not None:
                                            run.font.size = Pt(font_size)
                                        format_count += 1
                                    except Exception as e:
                                        self.logger.error(f"Error applying font style in table: {e}")
                                        return EditResult(
                                            success=False,
                                            error=f"Error applying font style in table: {str(e)}",
                                            details={"text": text, "font_name": font_name, "font_size": font_size, "error": str(e)}
                                        )
        
        if not text_found:
            return EditResult(
                success=False,
                error=f"Text '{text}' not found in document",
                details={"text": text, "font_name": font_name, "font_size": font_size}
            )
        
        if format_count == 0:
            # This can happen if we found the text but couldn't apply the font style
            return EditResult(
                success=False,
                error=f"Found '{text}' but could not apply font style",
                details={"text": text, "font_name": font_name, "font_size": font_size}
            )
        
        size_message = f" with size {font_size}pt" if font_size is not None else ""
        return EditResult(
            success=True,
            message=f"Applied font '{font_name}'{size_message} to '{text}' {format_count} times",
            count=format_count
        )
    
    def apply_alignment(self, doc: Document, text: str, alignment: str) -> EditResult:
        """
        Fixed version of _apply_alignment that handles case where paragraph is not found.
        
        Args:
            doc: The document object to edit
            text: Text to align
            alignment: Alignment type ("left", "center", "right", "justify")
            
        Returns:
            EditResult with operation status
        """
        if not text:
            return EditResult(
                success=False,
                error="Text to align must be provided"
            )
            
        # Map alignment string to WD_ALIGN_PARAGRAPH constants
        alignment_map = {
            "left": text_enums.WD_ALIGN_PARAGRAPH.LEFT,
            "center": text_enums.WD_ALIGN_PARAGRAPH.CENTER,
            "right": text_enums.WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": text_enums.WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        if alignment not in alignment_map:
            return EditResult(
                success=False,
                error=f"Unsupported alignment type: '{alignment}'. Supported types are: {', '.join(alignment_map.keys())}",
                details={"text": text, "alignment": alignment}
            )
            
        alignment_value = alignment_map[alignment]
        format_count = 0
        paragraph_found = False
        
        # Format paragraphs
        for para in doc.paragraphs:
            if text in para.text and text == para.text.strip():
                paragraph_found = True
                try:
                    para.alignment = alignment_value
                    format_count += 1
                except Exception as e:
                    self.logger.error(f