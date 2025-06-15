"""Utility to convert between different document formats."""
import os
import shutil
import tempfile
from typing import Optional, Tuple, List
import logging
from pathlib import Path

# Import docx for Word processing
import docx
from docx import Document

# Third-party converters, handle import errors gracefully
try:
    import pypandoc
    PANDOC_AVAILABLE = True
except ImportError:
    PANDOC_AVAILABLE = False

# COM automation for Windows (for handling .doc files)
try:
    import pythoncom
    import win32com.client
    COM_AVAILABLE = True
except ImportError:
    COM_AVAILABLE = False

logger = logging.getLogger("lawglance.document_converter")


class FormatNotSupportedError(Exception):
    """Exception raised when document format is not supported."""
    pass


class ConversionError(Exception):
    """Exception raised when document conversion fails."""
    pass


class DocumentConverter:
    """Utility to convert between different document formats."""
    
    def __init__(self):
        """Initialize the document converter."""
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        
        # Supported input formats
        self.input_formats = {
            '.docx': self._load_docx,
            '.txt': self._load_txt,
            '.md': self._load_md,
            '.html': self._load_html,
        }
        
        # Supported output formats
        self.output_formats = {
            '.docx': self._save_docx,
            '.txt': self._save_txt,
            '.md': self._save_md,
            '.html': self._save_html,
            '.pdf': self._save_pdf,
        }
        
        # Add .doc format if COM automation is available
        if COM_AVAILABLE:
            self.input_formats['.doc'] = self._load_doc
            self.output_formats['.doc'] = self._save_doc
        
        # Add additional formats if pandoc is available
        if PANDOC_AVAILABLE:
            additional_formats = ['.rst', '.epub', '.odt', '.rtf']
            for fmt in additional_formats:
                self.input_formats[fmt] = self._load_with_pandoc
                self.output_formats[fmt] = self._save_with_pandoc
    
    def convert_document(
        self,
        input_path: str,
        output_path: str,
        encoding: str = 'utf-8'
    ) -> bool:
        """
        Convert a document from one format to another.
        
        Args:
            input_path: Path to input document
            output_path: Path to output document
            encoding: Text encoding for text-based formats
            
        Returns:
            True if conversion was successful
            
        Raises:
            FormatNotSupportedError: If input or output format is not supported
            ConversionError: If conversion fails
            FileNotFoundError: If input file doesn't exist
        """
        # Validate input file exists
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        # Get file extensions
        input_ext = os.path.splitext(input_path)[1].lower()
        output_ext = os.path.splitext(output_path)[1].lower()
        
        # Check if formats are supported
        if input_ext not in self.input_formats:
            raise FormatNotSupportedError(f"Input format {input_ext} is not supported")
        
        if output_ext not in self.output_formats:
            raise FormatNotSupportedError(f"Output format {output_ext} is not supported")
        
        try:
            # If it's the same format, just copy the file
            if input_ext == output_ext:
                shutil.copy2(input_path, output_path)
                return True
            
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
            
            # Special case: If pandoc is available and both formats are supported by pandoc
            if PANDOC_AVAILABLE and input_ext != '.docx' and output_ext != '.docx':
                try:
                    output = pypandoc.convert_file(
                        input_path,
                        output_ext[1:],  # Remove the dot
                        outputfile=output_path
                    )
                    return True
                except Exception as e:
                    self.logger.warning(f"Pandoc conversion failed: {e}. Falling back to standard conversion.")
            
            # Standard conversion: Load -> Save
            content = self.input_formats[input_ext](input_path, encoding)
            self.output_formats[output_ext](content, output_path, encoding)
            
            return True
            
        except Exception as e:
            self.logger.error(f"Conversion error: {e}")
            raise ConversionError(f"Failed to convert {input_path} to {output_path}: {str(e)}")
    
    def get_supported_formats(self) -> Tuple[List[str], List[str]]:
        """
        Get lists of supported input and output formats.
        
        Returns:
            Tuple of (input_formats, output_formats)
        """
        return (
            sorted(list(self.input_formats.keys())),
            sorted(list(self.output_formats.keys()))
        )
    
    # Document loading methods
    def _load_docx(self, file_path: str, encoding: str) -> Document:
        """Load a .docx document."""
        return Document(file_path)
    
    def _load_doc(self, file_path: str, encoding: str) -> Document:
        """Load a .doc document by converting to .docx first."""
        pythoncom.CoInitialize()
        try:
            # Convert to docx using COM automation
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            # Create a temporary file for the docx version
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp:
                temp_path = temp.name
            
            # Open and save as docx
            doc = word.Documents.Open(os.path.abspath(file_path))
            doc.SaveAs2(temp_path, FileFormat=16)  # 16 = docx
            doc.Close()
            word.Quit()
            
            # Load the docx file
            result = Document(temp_path)
            
            # Clean up
            os.unlink(temp_path)
            return result
        finally:
            pythoncom.CoUninitialize()
    
    def _load_txt(self, file_path: str, encoding: str) -> str:
        """Load a text document."""
        with open(file_path, 'r', encoding=encoding) as f:
            return f.read()
    
    def _load_md(self, file_path: str, encoding: str) -> str:
        """Load a markdown document."""
        return self._load_txt(file_path, encoding)
    
    def _load_html(self, file_path: str, encoding: str) -> str:
        """Load an HTML document."""
        return self._load_txt(file_path, encoding)
    
    def _load_with_pandoc(self, file_path: str, encoding: str) -> str:
        """Load a document using pandoc."""
        if not PANDOC_AVAILABLE:
            raise FormatNotSupportedError("Pandoc is not available")
        
        # Convert to markdown as intermediate format
        return pypandoc.convert_file(file_path, 'md')
    
    # Document saving methods
    def _save_docx(self, content: Any, file_path: str, encoding: str) -> None:
        """Save content as a .docx document."""
        if isinstance(content, Document):
            content.save(file_path)
        else:
            # Convert string content to docx
            doc = Document()
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            doc.save(file_path)
    
    def _save_doc(self, content: Any, file_path: str, encoding: str) -> None:
        """Save content as a .doc document."""
        if not COM_AVAILABLE:
            raise FormatNotSupportedError(".doc format requires COM automation")
        
        # First save as .docx
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp:
            temp_path = temp.name
            
        self._save_docx(content, temp_path, encoding)
        
        # Convert to .doc
        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(temp_path)
            doc.SaveAs2(file_path, FileFormat=0)  # 0 = .doc
            doc.Close()
            word.Quit()
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            pythoncom.CoUninitialize()
    
    def _save_txt(self, content: Any, file_path: str, encoding: str) -> None:
        """Save content as a text document."""
        if isinstance(content, str):
            with open(file_path, 'w', encoding=encoding) as f:
                f.write(content)
        elif isinstance(content, Document):
            # Extract text from docx
            text = "\n\n".join([p.text for p in content.paragraphs if p.text.strip()])
            with open(file_path, 'w', encoding=encoding) as f:
                f.write(text)
    
    def _save_md(self, content: Any, file_path: str, encoding: str) -> None:
        """Save content as a markdown document."""
        self._save_txt(content, file_path, encoding)
    
    def _save_html(self, content: Any, file_path: str, encoding: str) -> None:
        """Save content as an HTML document."""
        if isinstance(content, str):
            # Very basic markdown to HTML conversion
            html_content = f"<!DOCTYPE html>\n<html>\n<head>\n<title>Converted Document</title>\n</head>\n<body>\n"
            
            paragraphs = content.split('\n\n')
            for p in paragraphs:
                if p.strip():
                    html_content += f"<p>{p}</p>\n"
            
            html_content += "</body>\n</html>"
            
            with open(file_path, 'w', encoding=encoding) as f:
                f.write(html_content)
        elif isinstance(content, Document):
            # Convert docx content to HTML
            html_content = f"<!DOCTYPE html>\n<html>\n<head>\n<title>Converted Document</title>\n</head>\n<body>\n"
            
            for p in content.paragraphs:
                if p.text.strip():
                    html_content += f"<p>{p.text}</p>\n"
            
            html_content += "</body>\n</html>"
            
            with open(file_path, 'w', encoding=encoding) as f:
                f.write(html_content)
    
    def _save_pdf(self, content: Any, file_path: str, encoding: str) -> None:
        """Save content as a PDF document."""
        if not COM_AVAILABLE:
            raise FormatNotSupportedError("PDF export requires COM automation")
        
        # First save as .docx
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp:
            temp_path = temp.name
            
        self._save_docx(content, temp_path, encoding)
        
        # Convert to PDF
        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            doc = word.Documents.Open(temp_path)
            doc.SaveAs2(file_path, FileFormat=17)  # 17 = PDF
            doc.Close()
            word.Quit()
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            pythoncom.CoUninitialize()
    
    def _save_with_pandoc(self, content: Any, file_path: str, encoding: str) -> None:
        """Save content using pandoc."""
        if not PANDOC_AVAILABLE:
            raise FormatNotSupportedError("Pandoc is not available")
        
        # We need to save the content to a temporary file first
        with tempfile.NamedTemporaryFile(suffix='.md', mode='w', encoding=encoding, delete=False) as temp:
            if isinstance(content, str):
                temp.write(content)
            elif isinstance(content, Document):
                temp.write("\n\n".join([p.text for p in content.paragraphs if p.text.strip()]))
            temp_path = temp.name
        
        try:
            # Get the output format from the file extension
            output_format = os.path.splitext(file_path)[1][1:]  # Remove the dot
            
            # Convert using pandoc
            pypandoc.convert_file(
                temp_path,
                output_format,
                outputfile=file_path
            )
        finally:
            # Clean up the temporary file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
