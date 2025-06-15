"""Improved document comparison utility with dataclasses."""
import os
import re
import difflib
import logging
from typing import Dict, List, Tuple, Optional, Any, Union, Set
from dataclasses import dataclass, field
from pathlib import Path

import docx
from docx import Document

logger = logging.getLogger("lawglance.document_compare")

@dataclass
class ComparisonResult:
    """Container for document comparison results."""
    success: bool = True
    similarity_percentage: float = 0.0
    difference_count: int = 0
    additions: List[str] = field(default_factory=list)
    deletions: List[str] = field(default_factory=list)
    changes: List[Dict[str, str]] = field(default_factory=list)
    structure_differences: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None
    summary: str = ""

    def __repr__(self) -> str:
        """Return string representation of comparison results."""
        if not self.success:
            return f"ComparisonResult(success=False, error='{self.error}')"
        return (f"ComparisonResult(similarity={self.similarity_percentage:.1f}%, "
                f"differences={self.difference_count})")
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for API responses."""
        return {
            "success": self.success,
            "similarity_percentage": round(self.similarity_percentage, 2),
            "difference_count": self.difference_count,
            "additions": self.additions,
            "deletions": self.deletions,
            "changes": self.changes,
            "structure_differences": self.structure_differences,
            "summary": self.summary,
            "error": self.error
        }

@dataclass
class DocumentInfo:
    """Container for document metadata and content."""
    path: str
    content: str = ""
    paragraphs: List[str] = field(default_factory=list)
    sections: List[str] = field(default_factory=list)
    word_count: int = 0
    
    def __post_init__(self):
        """Initialize after dataclass initialization."""
        if not self.content and os.path.exists(self.path):
            self.load_content()
            
    def load_content(self) -> None:
        """Load document content from file."""
        file_ext = os.path.splitext(self.path)[1].lower()
        
        try:
            if file_ext == ".docx":
                doc = Document(self.path)
                self.paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                self.content = "\n".join(self.paragraphs)
                
            elif file_ext in [".txt", ".md"]:
                with open(self.path, 'r', encoding='utf-8') as f:
                    self.content = f.read()
                self.paragraphs = [p for p in self.content.split('\n\n') if p.strip()]
                
            # Extract sections and calculate word count
            self.sections = self._extract_sections()
            self.word_count = len(self.content.split())
                
        except Exception as e:
            logger.error(f"Error loading document {self.path}: {str(e)}")
            raise
            
    def _extract_sections(self) -> List[str]:
        """Extract section titles from content."""
        # Common patterns for section headers in legal documents
        patterns = [
            r'^([IVX]+\.\s.+)$',                     # Roman numerals: IV. SECTION
            r'^([0-9]+\.\s.+)$',                     # Numbered: 1. SECTION
            r'^((?:[A-Z][A-Z\s]+))$',                # ALL CAPS SECTION
            r'^([A-Z][a-z].+:)',                     # Title case ending with colon: Section:
            r'^(ARTICLE\s+[IVX0-9]+\s*[-—:]\s*.+)$'  # ARTICLE IV - SECTION
        ]
        
        sections = []
        for line in self.content.splitlines():
            line = line.strip()
            for pattern in patterns:
                match = re.match(pattern, line, re.MULTILINE)
                if match:
                    sections.append(match.group(1))
                    break
        
        return sections


@dataclass
class DocumentCompare:
    """Compares legal documents and identifies differences."""
    file1_path: Optional[str] = None
    file2_path: Optional[str] = None
    doc1: Optional[DocumentInfo] = None
    doc2: Optional[DocumentInfo] = None
    supported_formats: List[str] = field(default_factory=lambda: [".txt", ".docx", ".md"])
    
    def __post_init__(self):
        """Initialize after object creation."""
        # If paths are provided but docs aren't, create DocumentInfo objects
        if self.file1_path and not self.doc1:
            if os.path.exists(self.file1_path):
                self.doc1 = DocumentInfo(self.file1_path)
            
        if self.file2_path and not self.doc2:
            if os.path.exists(self.file2_path):
                self.doc2 = DocumentInfo(self.file2_path)
    
    def __repr__(self) -> str:
        """Return string representation of the document comparer."""
        return f"DocumentCompare(file1={self.file1_path or 'None'}, file2={self.file2_path or 'None'})"
    
    def compare_files(self, file1_path: Optional[str] = None, file2_path: Optional[str] = None) -> ComparisonResult:
        """
        Compare two documents and identify differences.
        
        Args:
            file1_path: Path to the first document (optional if set in constructor)
            file2_path: Path to the second document (optional if set in constructor)
            
        Returns:
            ComparisonResult containing comparison results
        """
        # Use provided paths or fall back to instance variables
        path1 = file1_path or self.file1_path
        path2 = file2_path or self.file2_path
        
        if not path1 or not path2:
            return ComparisonResult(
                success=False,
                error="Both file paths are required for comparison"
            )
            
        try:
            # Validate files exist
            for path in [path1, path2]:
                if not os.path.exists(path):
                    return ComparisonResult(
                        success=False, 
                        error=f"File not found: {path}"
                    )
            
            # Validate file formats
            for path in [path1, path2]:
                ext = os.path.splitext(path)[1].lower()
                if ext not in self.supported_formats:
                    return ComparisonResult(
                        success=False, 
                        error=f"Unsupported file format: {ext}"
                    )
            
            # Load documents
            doc1 = DocumentInfo(path1)
            doc2 = DocumentInfo(path2)
            
            # Save for future reference
            self.doc1 = doc1
            self.doc2 = doc2
            self.file1_path = path1
            self.file2_path = path2
            
            return self.compare_text(doc1.content, doc2.content)
            
        except Exception as e:
            logger.error(f"Error comparing files: {str(e)}")
            return ComparisonResult(success=False, error=str(e))
    
    def compare_text(self, text1: str, text2: str) -> ComparisonResult:
        """
        Compare two text strings and identify differences.
        
        Args:
            text1: First text string
            text2: Second text string
            
        Returns:
            ComparisonResult containing comparison results
        """
        # Split texts into lines
        lines1 = text1.splitlines()
        lines2 = text2.splitlines()
        
        # Calculate similarity ratio
        similarity = difflib.SequenceMatcher(None, text1, text2).ratio()
        
        # Find differences
        differ = difflib.Differ()
        diff = list(differ.compare(lines1, lines2))
        
        # Organize differences
        additions = [line[2:] for line in diff if line.startswith('+ ')]
        deletions = [line[2:] for line in diff if line.startswith('- ')]
        
        # Find changed sections using unified_diff for context
        diff_with_context = list(difflib.unified_diff(
            lines1, lines2, lineterm='', n=3
        ))[3:]  # Skip the headers
        
        changes = []
        for i, line in enumerate(diff_with_context):
            if line.startswith('+') and i > 0 and diff_with_context[i-1].startswith('-'):
                changes.append({
                    "removed": diff_with_context[i-1][1:].strip(),
                    "added": line[1:].strip()
                })
        
        # Count differences
        diff_count = len(additions) + len(deletions)
        
        # Extract structural differences (sections, paragraphs)
        structure_diff = self._compare_structure(text1, text2)
        
        # Create the result
        result = ComparisonResult(
            success=True,
            similarity_percentage=similarity * 100,
            difference_count=diff_count,
            additions=additions,
            deletions=deletions,
            changes=changes,
            structure_differences=structure_diff
        )
        
        # Generate summary
        result.summary = self._generate_summary(
            similarity, 
            additions, 
            deletions, 
            changes, 
            structure_diff
        )
        
        return result
    
    def find_common_terms(self) -> Set[str]:
        """Find common legal terms in both documents."""
        if not self.doc1 or not self.doc2:
            return set()
            
        # Define patterns for legal terms
        legal_term_pattern = re.compile(
            r'\b(plaintiff|defendant|court|judgment|contract|breach|damages|'
            r'injunction|appeal|petition|statute|regulation|clause|jurisdiction|'
            r'tort|liability|settlement|arbitration|hearing|testimony|evidence)\b',
            re.IGNORECASE
        )
        
        # Find terms in both documents
        terms1 = set(legal_term_pattern.findall(self.doc1.content.lower()))
        terms2 = set(legal_term_pattern.findall(self.doc2.content.lower()))
        
        # Return intersection
        return terms1.intersection(terms2)
    
    @staticmethod
    def _compare_structure(text1: str, text2: str) -> Dict[str, Any]:
        """
        Compare the structure of two documents.
        
        Args:
            text1: First text string
            text2: Second text string
            
        Returns:
            Dictionary with structural differences
        """
        # Extract sections using common patterns in legal documents
        sections1 = DocumentCompare._extract_sections_from_text(text1)
        sections2 = DocumentCompare._extract_sections_from_text(text2)
        
        # Find sections present in one document but not the other
        unique_sections1 = [s for s in sections1 if s not in sections2]
        unique_sections2 = [s for s in sections2 if s not in sections1]
        
        # Count paragraphs
        paragraphs1 = [p for p in text1.split('\n\n') if p.strip()]
        paragraphs2 = [p for p in text2.split('\n\n') if p.strip()]
        
        return {
            "sections_only_in_first": unique_sections1,
            "sections_only_in_second": unique_sections2,
            "paragraph_count_first": len(paragraphs1),
            "paragraph_count_second": len(paragraphs2),
        }
    
    @staticmethod
    def _extract_sections_from_text(text: str) -> List[str]:
        """
        Extract section titles from text.
        
        Args:
            text: Text to extract sections from
            
        Returns:
            List of section titles
        """
        # Common patterns for section headers in legal documents
        patterns = [
            r'^([IVX]+\.\s.+)$',                     # Roman numerals: IV. SECTION
            r'^([0-9]+\.\s.+)$',                     # Numbered: 1. SECTION
            r'^((?:[A-Z][A-Z\s]+))$',                # ALL CAPS SECTION
            r'^([A-Z][a-z].+:)',                     # Title case ending with colon: Section:
            r'^(ARTICLE\s+[IVX0-9]+\s*[-—:]\s*.+)$'  # ARTICLE IV - SECTION
        ]
        
        sections = []
        for line in text.splitlines():
            line = line.strip()
            for pattern in patterns:
                match = re.match(pattern, line, re.MULTILINE)
                if match:
                    sections.append(match.group(1))
                    break
        
        return sections
    
    @staticmethod
    def _generate_summary(similarity: float, additions: List[str], deletions: List[str], 
                          changes: List[Dict[str, str]], structure_diff: Dict[str, Any]) -> str:
        """
        Generate a human-readable summary of the differences.
        
        Args:
            similarity: Similarity ratio between documents
            additions: List of added lines
            deletions: List of deleted lines
            changes: List of changed lines
            structure_diff: Dictionary with structural differences
            
        Returns:
            Summary text
        """
        summary = []
        
        # Overall similarity
        if similarity > 0.95:
            summary.append("The documents are nearly identical.")
        elif similarity > 0.8:
            summary.append("The documents are very similar with some differences.")
        elif similarity > 0.6:
            summary.append("The documents have significant differences but share much content.")
        elif similarity > 0.4:
            summary.append("The documents are more different than similar.")
        else:
            summary.append("The documents are substantially different.")
            
        # Additions and deletions
        if additions:
            if len(additions) == 1:
                summary.append("There is 1 added line.")
            else:
                summary.append(f"There are {len(additions)} added lines.")
                
        if deletions:
            if len(deletions) == 1:
                summary.append("There is 1 deleted line.")
            else:
                summary.append(f"There are {len(deletions)} deleted lines.")
        
        # Structure differences
        sections_only_in_first = structure_diff.get("sections_only_in_first", [])
        sections_only_in_second = structure_diff.get("sections_only_in_second", [])
        
        if sections_only_in_first:
            if len(sections_only_in_first) == 1:
                summary.append(f"The section '{sections_only_in_first[0]}' is only in the first document.")
            else:
                summary.append(f"{len(sections_only_in_first)} sections are only in the first document.")
                
        if sections_only_in_second:
            if len(sections_only_in_second) == 1:
                summary.append(f"The section '{sections_only_in_second[0]}' is only in the second document.")
            else:
                summary.append(f"{len(sections_only_in_second)} sections are only in the second document.")
        
        # Paragraph count differences
        para_first = structure_diff.get("paragraph_count_first", 0)
        para_second = structure_diff.get("paragraph_count_second", 0)
        
        if para_first != para_second:
            summary.append(f"The first document has {para_first} paragraphs, while the second has {para_second}.")
        
        return " ".join(summary)
