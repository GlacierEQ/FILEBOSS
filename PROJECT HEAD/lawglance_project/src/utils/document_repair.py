"""
Document repair and scanning utilities for Word documents.
Provides tools to identify and fix common document issues.
"""
import logging
import re
from typing import Dict, List, Tuple, Optional, Any, Set
from dataclasses import dataclass, field
from enum import Enum
from docx import Document
from docx.shared import Pt
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

# Configure logger
logger = logging.getLogger("lawglance.document_repair")


class IssueType(Enum):
    """Types of document issues that can be detected."""
    FORMAT_INCONSISTENCY = "Format Inconsistency"
    STYLE_INCONSISTENCY = "Style Inconsistency" 
    HEADING_HIERARCHY = "Heading Hierarchy Issue"
    ORPHANED_PARAGRAPH = "Orphaned Paragraph"
    TABLE_ISSUE = "Table Structure Issue"
    WHITESPACE_REDUNDANCY = "Whitespace Redundancy"
    EMPTY_SECTION = "Empty Section"
    DUPLICATE_CONTENT = "Duplicate Content"
    MISSING_REFERENCE = "Missing Reference"
    BROKEN_LINK = "Broken Link"
    INCONSISTENT_NUMBERING = "Inconsistent Numbering"


@dataclass
class DocumentIssue:
    """Represents an issue found in a document."""
    issue_type: IssueType
    description: str
    location: str  # Description of where the issue is located
    element_index: Optional[int] = None  # Index of the paragraph/table where issue occurs
    can_auto_fix: bool = False
    severity: int = 1  # 1-3, where 3 is most severe
    
    def __str__(self) -> str:
        """String representation of the issue."""
        return f"{self.issue_type.value} at {self.location}: {self.description}"


@dataclass
class RepairReport:
    """Report of document issues and fixes."""
    issues_found: List[DocumentIssue] = field(default_factory=list)
    issues_fixed: List[DocumentIssue] = field(default_factory=list)
    scan_complete: bool = False
    
    def add_issue(self, issue: DocumentIssue) -> None:
        """Add an issue to the report."""
        self.issues_found.append(issue)
    
    def mark_fixed(self, issue: DocumentIssue) -> None:
        """Mark an issue as fixed."""
        if issue not in self.issues_fixed:
            self.issues_fixed.append(issue)
    
    @property
    def total_issues(self) -> int:
        """Get the total number of issues found."""
        return len(self.issues_found)
    
    @property
    def fixed_count(self) -> int:
        """Get the number of issues fixed."""
        return len(self.issues_fixed)
    
    @property
    def unfixed_count(self) -> int:
        """Get the number of issues not fixed."""
        return self.total_issues - self.fixed_count
    
    @property
    def summary(self) -> Dict[str, Any]:
        """Get a summary of the report."""
        issues_by_type = {}
        for issue in self.issues_found:
            issue_type = issue.issue_type.value
            if issue_type not in issues_by_type:
                issues_by_type[issue_type] = {"found": 0, "fixed": 0}
            issues_by_type[issue_type]["found"] += 1
            
        for issue in self.issues_fixed:
            issue_type = issue.issue_type.value
            issues_by_type[issue_type]["fixed"] += 1
            
        return {
            "total_found": self.total_issues,
            "total_fixed": self.fixed_count,
            "by_type": issues_by_type
        }
    
    def __str__(self) -> str:
        """String representation of the report."""
        result = f"Document Repair Report: {self.total_issues} issues found, {self.fixed_count} fixed\n"
        if self.total_issues > 0:
            result += "\nIssues by type:\n"
            for issue_type, counts in self.summary["by_type"].items():
                result += f"- {issue_type}: {counts['found']} found, {counts['fixed']} fixed\n"
                
        return result


class DocumentRepair:
    """Utilities for scanning and repairing Word documents."""
    
    def __init__(self, auto_fix: bool = False):
        """
        Initialize the document repair tool.
        
        Args:
            auto_fix: Whether to automatically fix issues when possible
        """
        self.auto_fix = auto_fix
        self.report = RepairReport()
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
    
    def scan_document(self, doc: Document) -> RepairReport:
        """
        Scan a document for issues.
        
        Args:
            doc: Document to scan
            
        Returns:
            Report of issues found
        """
        self.report = RepairReport()
        
        # Run different scanners
        self._scan_heading_hierarchy(doc)
        self._scan_formatting_consistency(doc)
        self._scan_whitespace_issues(doc)
        self._scan_table_issues(doc)
        self._scan_empty_sections(doc)
        self._scan_numbering_consistency(doc)
        self._scan_duplicate_content(doc)
        
        # Mark scan as complete
        self.report.scan_complete = True
        return self.report
    
    def repair_document(self, doc: Document) -> RepairReport:
        """
        Repair issues in a document.
        
        Args:
            doc: Document to repair
            
        Returns:
            Report of issues fixed
        """
        # Scan if not already done
        if not self.report.scan_complete:
            self.scan_document(doc)
        
        # Try to fix each issue
        for issue in self.report.issues_found:
            if not issue.can_auto_fix:
                continue
                
            try:
                fixed = self._fix_issue(doc, issue)
                if fixed:
                    self.report.mark_fixed(issue)
            except ValueError as e:
                self.logger.warning("Failed to fix issue: %s. Error: %s", issue, e)
        
        return self.report
    
    def _scan_heading_hierarchy(self, doc: Document) -> None:
        """
        Scan for heading hierarchy issues.
        
        Args:
            doc: Document to scan
        """
        # Track heading levels
        # Removed unused variable current_level
        
        # Collect all headings and their levels
        headings = []
        for i, para in enumerate(doc.paragraphs):
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading', ''))
                    headings.append((i, level, para.text))
                except ValueError:
                    # If we can't parse the heading level, skip it
                    continue
        
        # Check for hierarchy issues (e.g., Heading1 -> Heading3 without Heading2)
        for i in range(1, len(headings)):
            prev_idx, prev_level, prev_text = headings[i-1]
            curr_idx, curr_level, curr_text = headings[i]
            
            # Check if heading level jumps by more than 1
            if curr_level > prev_level + 1:
                self.report.add_issue(DocumentIssue(
                    issue_type=IssueType.HEADING_HIERARCHY,
                    description=f"Heading level jumps from {prev_level} to {curr_level} without intermediate levels",
                    location=f"Paragraph {curr_idx} ('{curr_text}')",
                    element_index=curr_idx,
                    can_auto_fix=True,
                    severity=2
                ))
    
    def _scan_formatting_consistency(self, doc: Document) -> None:
        """
        Scan for formatting consistency issues.
        
        Args:
            doc: Document to scan
        """
        # Track font usage by style to detect inconsistencies
        style_fonts = {}
        style_sizes = {}
        
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name
            
            # Skip empty paragraphs
            if not para.text.strip():
                continue
            
            for run in para.runs:
                if not run.text.strip():
                    continue
                
                # Check font consistency
                font_name = run.font.name
                if font_name:
                    if style_name not in style_fonts:
                        style_fonts[style_name] = {font_name: 1}
                    elif font_name not in style_fonts[style_name]:
                        style_fonts[style_name][font_name] = 1
                    else:
                        style_fonts[style_name][font_name] += 1
                
                # Check font size consistency
                font_size = run.font.size
                if font_size:
                    if style_name not in style_sizes:
                        style_sizes[style_name] = {font_size: 1}
                    elif font_size not in style_sizes[style_name]:
                        style_sizes[style_name][font_size] = 1
                    else:
                        style_sizes[style_name][font_size] += 1
        
        # Check for inconsistent fonts within styles
        for style_name, fonts in style_fonts.items():
            if len(fonts) > 1:
                # Find the most common font
                most_common_font = max(fonts.items(), key=lambda x: x[1])[0]
                other_fonts = [f for f in fonts.keys() if f != most_common_font]
                
                self.report.add_issue(DocumentIssue(
                    issue_type=IssueType.FORMAT_INCONSISTENCY,
                    description=f"Multiple fonts used with '{style_name}' style. " +
                               f"Main font: {most_common_font}, Others: {', '.join(other_fonts)}",
                    location=f"Throughout document",
                    can_auto_fix=True,
                    severity=1
                ))
        
        # Check for inconsistent font sizes within styles
        for style_name, sizes in style_sizes.items():
            if len(sizes) > 1:
                # Find the most common size
                most_common_size = max(sizes.items(), key=lambda x: x[1])[0]
                other_sizes = [str(s) for s in sizes.keys() if s != most_common_size]
                
                self.report.add_issue(DocumentIssue(
                    issue_type=IssueType.FORMAT_INCONSISTENCY,
                    description=f"Multiple font sizes used with '{style_name}' style. " +
                               f"Main size: {most_common_size}, Others: {', '.join(other_sizes)}",
                    location=f"Throughout document",
                    can_auto_fix=True,
                    severity=1
                ))
    
    def _scan_whitespace_issues(self, doc: Document) -> None:
        """
        Scan for whitespace issues like double spaces or excessive line breaks.
        
        Args:
            doc: Document to scan
        """
        for i, para in enumerate(doc.paragraphs):
            # Check for double spaces
            if "  " in para.text:
                self.report.add_issue(DocumentIssue(
                    issue_type=IssueType.WHITESPACE_REDUNDANCY,
                    description="Multiple consecutive spaces detected",
                    location=f"Paragraph {i}: '{para.text[:30]}...'",
                    element_index=i,
                    can_auto_fix=True,
                    severity=1
                ))
            
            # Check for excessive line breaks (empty paragraphs)
            if i > 0 and i < len(doc.paragraphs) - 1:
                if not para.text.strip() and not doc.paragraphs[i-1].text.strip():
                    self.report.add_issue(DocumentIssue(
                        issue_type=IssueType.WHITESPACE_REDUNDANCY,
                        description="Multiple consecutive empty paragraphs",
                        location=f"Paragraph {i}",
                        element_index=i,
                        can_auto_fix=True,
                        severity=1
                    ))
    
    def _scan_table_issues(self, doc: Document) -> None:
        """
        Scan for issues in tables.
        
        Args:
            doc: Document to scan
        """
        for i, table in enumerate(doc.tables):
            # Check for empty cells
            empty_cells = 0
            total_cells = 0
            
            for row in table.rows:
                for cell in row.cells:
                    total_cells += 1
                    if not cell.text.strip():
                        empty_cells += 1
            
            # If more than half the cells are empty, flag it
            if empty_cells > total_cells / 2:
                self.report.add_issue(DocumentIssue(
                    issue_type=IssueType.TABLE_ISSUE,
                    description=f"Table has many empty cells ({empty_cells} of {total_cells})",
                    location=f"Table {i}",
                    element_index=i,
                    can_auto_fix=False,
                    severity=1
                ))
            
            # Check for inconsistent row counts (jagged tables)
            row_lengths = {}
            for j, row in enumerate(table.rows):
                row_length = len(row.cells)
                if row_length not in row_lengths:
                    row_lengths[row_length] = []
                row_lengths[row_length].append(j)
                
            if len(row_lengths) > 1:
                description = "Table has inconsistent row lengths: " + \
                              ", ".join([f"{length} cells ({len(rows)} rows)" 
                                       for length, rows in row_lengths.items()])
                
                self.report.add_issue(DocumentIssue(
                    issue_type=IssueType.TABLE_ISSUE,
                    description=description,
                    location=f"Table {i}",
                    element_index=i,
                    can_auto_fix=False,
                    severity=2
                ))
    
    def _scan_empty_sections(self, doc: Document) -> None:
        """
        Scan for empty sections (headings with no content).
        
        Args:
            doc: Document to scan
        """
        for i in range(len(doc.paragraphs) - 1):
            para = doc.paragraphs[i]
            next_para = doc.paragraphs[i + 1]
            
            if para.style.name.startswith('Heading'):
                # Check if the next paragraph is also a heading or empty
                if next_para.style.name.startswith('Heading') or not next_para.text.strip():
                    self.report.add_issue(DocumentIssue(
                        issue_type=IssueType.EMPTY_SECTION,
                        description=f"Heading '{para.text}' has no content",
                        location=f"Paragraph {i}",
                        element_index=i,
                        can_auto_fix=False,
                        severity=2
                    ))
    
    def _scan_numbering_consistency(self, doc: Document) -> None:
        """
        Scan for inconsistent numbering in lists and sections.
        
        Args:
            doc: Document to scan
        """
        # This is a simplified version - a real implementation would need to
        # handle Microsoft Word's complex numbering schemes
        
        # Look for numbered paragraphs with patterns like "1.", "2.", etc.
        numbered_paragraphs = []
        
        for i, para in enumerate(doc.paragraphs):
            # Check for common numbering patterns
            match = re.match(r'^(\d+)\.?\s', para.text)
            if match:
                number = int(match.group(1))
                numbered_paragraphs.append((i, number, para.text))
        
        # Check for sequence breaks
        for j in range(1, len(numbered_paragraphs)):
            prev_idx, prev_num = numbered_paragraphs[j-1]
            curr_idx, curr_num, curr_text = numbered_paragraphs[j]
            
            # If there are paragraphs in between, this might be a new list
            if curr_idx - prev_idx > 1:
                # Check if the current number is 1, suggesting a new list
                if curr_num != 1:
                    self.report.add_issue(DocumentIssue(
                        issue_type=IssueType.INCONSISTENT_NUMBERING,
                        description="New numbered list doesn't start with 1",
                        location=f"Paragraph {curr_idx}: '{curr_text[:30]}...'",
                        element_index=curr_idx,
                        can_auto_fix=True,
                        severity=1
                    ))
            else:
                # Check sequence (should increment by 1)
                if curr_num != prev_num + 1:
                    self.report.add_issue(DocumentIssue(
                        issue_type=IssueType.INCONSISTENT_NUMBERING,
                        description=f"Number sequence goes from {prev_num} to {curr_num}",
                        location=f"Paragraph {curr_idx}: '{curr_text[:30]}...'",
                        element_index=curr_idx,
                        can_auto_fix=True,
                        severity=2
                    ))
    
    def _scan_duplicate_content(self, doc: Document) -> None:
        """
        Scan for duplicate paragraphs.
        
        Args:
            doc: Document to scan
        """
        content_map = {}
        
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip() or len(para.text) < 20:
                continue  # Skip empty or very short paragraphs
                
            text = para.text.strip()
            if text not in content_map:
                content_map[text] = [i]
            else:
                content_map[text].append(i)
        
        # Report duplicates
        for text, positions in content_map.items():
            if len(positions) > 1:
                self.report.add_issue(DocumentIssue(
                    issue_type=IssueType.DUPLICATE_CONTENT,
                    description=f"Duplicate paragraph found {len(positions)} times",
                    location=f"Paragraphs {', '.join(map(str, positions))}",
                    element_index=positions[0],  # Reference the first occurrence
                    can_auto_fix=True,
                    severity=2
                ))
    
    def _fix_issue(self, doc: Document, issue: DocumentIssue) -> bool:
        """
        Attempt to fix a document issue.
        
        Args:
            doc: Document to fix
            issue: Issue to fix
            
        Returns:
            True if the issue was fixed, False otherwise
        """
        if issue.issue_type == IssueType.WHITESPACE_REDUNDANCY:
            return self._fix_whitespace(doc, issue)
            
        elif issue.issue_type == IssueType.FORMAT_INCONSISTENCY:
            return self._fix_formatting(doc, issue)
            
        elif issue.issue_type == IssueType.HEADING_HIERARCHY:
            return self._fix_heading_hierarchy(doc, issue)
            
        elif issue.issue_type == IssueType.INCONSISTENT_NUMBERING:
            return self._fix_numbering(doc, issue)
            
        elif issue.issue_type == IssueType.DUPLICATE_CONTENT:
            return self._fix_duplicates(doc, issue)
            
        # For other issue types, no automatic fix is available yet
        return False
    
    def _fix_whitespace(self, doc: Document, issue: DocumentIssue) -> bool:
        """Fix whitespace issues."""
        if issue.element_index is None:
            return False
        
        para = doc.paragraphs[issue.element_index]
        
        # Fix double spaces
        if "  " in para.text:
            for run in para.runs:
                # Keep replacing double spaces until there are none left
                while "  " in run.text:
                    run.text = run.text.replace("  ", " ")
            return True
        
        # Fix consecutive empty paragraphs
        if not para.text.strip():
            # Check if previous paragraph is also empty
            if (issue.element_index > 0 and 
                not doc.paragraphs[issue.element_index - 1].text.strip()):
                # This gets complex with Word's XML structure
                # In a real implementation, we'd need to modify the 
                # document's XML to remove the paragraph element
                self.logger.info("Removing empty paragraph would require XML manipulation")
                return False
        
        return False
    
    def _fix_formatting(self, doc: Document, issue: DocumentIssue) -> bool:
        """Fix formatting inconsistencies."""
        description = issue.description.lower()
        
        # Implement logic to fix formatting inconsistencies
        style_fonts = {}
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name != most_common_font:
                    run.font.name = most_common_font
                    self.logger.info(f"Fixed font inconsistency in paragraph: '{para.text}'")
        return True
        
        if "font" in description:
            self.logger.info("Would fix font inconsistencies: %s", issue.description)
            return True
            
        if "size" in description:
            self.logger.info("Would fix font size inconsistencies: %s", issue.description)
            return True
        
        return False
    
    def _fix_heading_hierarchy(self, doc: Document, issue: DocumentIssue) -> bool:
        """Fix heading hierarchy issues."""
        if issue.element_index is None:
            return False
        
        para = doc.paragraphs[issue.element_index]
        
        # Attempt to determine the correct heading level
        if para.style.name.startswith('Heading'):
            try:
                current_level = int(para.style.name.replace('Heading', ''))
                # Find previous heading's level
                prev_level = 0
                for i in range(issue.element_index - 1, -1, -1):
                    prev_para = doc.paragraphs[i]
                    if prev_para.style.name.startswith('Heading'):
                        prev_level = int(prev_para.style.name.replace('Heading', ''))
                        break
                
                # If we found a gap, adjust the level
                if current_level > prev_level + 1:
                    correct_level = prev_level + 1
                    para.style = f'Heading{correct_level}'
                    return True
            except (ValueError, IndexError):
                return False
        
        return False
    
    def _fix_numbering(self, doc: Document, issue: DocumentIssue) -> bool:
        """Fix numbering inconsistencies."""
        if issue.element_index is None:
            return False
        
        para = doc.paragraphs[issue.element_index]
        text = para.text
        
        # Extract the current number
        match = re.match(r'^(\d+)\.?\s+(.*)$', text)
        if not match:
            return False
            
        current_num = int(match.group(1))
        rest_of_text = match.group(2)
        
        # Determine expected number
        expected_num = 1  # Default for new lists
        
        # Check if this is in a sequence
        for i in range(issue.element_index - 1, -1, -1):
            # Removed unused variable prev_text
            prev_match = re.match(r'^(\d+)\.?\s', prev_text)
            if prev_match:
                expected_num = int(prev_match.group(1)) + 1
                break
            elif i < issue.element_index - 1:
                # If we've skipped some paragraphs, assume it's a new list
                expected_num = 1
                break
        
        # Apply the correction
        if current_num != expected_num:
            # Replace the text with the correct number
            new_text = f"{expected_num}. {rest_of_text}"
            
            # Update runs (simplistic approach - real implementation would be more careful)
            if para.runs:
                para.runs[0].text = new_text
                for i in range(1, len(para.runs)):
                    para.runs[i].text = ""
            else:
                para.text = new_text
                
            return True
        
        return False
    
    def _fix_duplicates(self, doc: Document, issue: DocumentIssue) -> bool:
        """
        Fix duplicate content by marking duplicates.
        
        Note: Actually removing duplicates is complex with python-docx
        as it requires XML manipulation.
        """
        if issue.element_index is None:
            return False
        
        # Find all paragraphs with identical text
        target_text = doc.paragraphs[issue.element_index].text
        duplicate_indices = []
        
        for i, para in enumerate(doc.paragraphs):
            if para.text == target_text and i != issue.element_index:
                duplicate_indices.append(i)
        
        # Mark duplicates (real implementation might delete them)
        for idx in duplicate_indices:
            para = doc.paragraphs[idx]
            if para.runs:
                para.runs[0].font.highlight_color = 7  # Yellow
                
                # Add note at the beginning
                original_text = para.runs[0].text
                para.runs[0].text = "[DUPLICATE] " + original_text
        
        return len(duplicate_indices) > 0
