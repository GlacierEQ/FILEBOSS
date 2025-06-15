"""Performance optimizations for document processing."""
import os
import re
import time
from typing import List, Dict, Any, Callable, Optional, Set, Tuple, Union
import functools
import logging
from docx import Document

logger = logging.getLogger("lawglance.performance")

# Performance monitoring decorator
def measure_performance(func):
    """Decorator to measure execution time of functions."""
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        execution_time = end_time - start_time
        logger.debug(f"Function {func.__name__} took {execution_time:.4f} seconds to execute")
        return result
    return wrapper

class TextCache:
    """Cache for text operations to avoid redundant processing."""
    
    def __init__(self, max_size: int = 100):
        """Initialize text cache with maximum size."""
        self.cache = {}
        self.max_size = max_size
        self.hits = 0
        self.misses = 0
    
    def get(self, key: str) -> Optional[Any]:
        """Get a value from cache."""
        if key in self.cache:
            self.hits += 1
            return self.cache[key]
        self.misses += 1
        return None
    
    def set(self, key: str, value: Any) -> None:
        """Set a value in cache."""
        # Evict an entry if we're at max capacity
        if len(self.cache) >= self.max_size and key not in self.cache:
            # Simple least-recently-used strategy: remove first item
            self.cache.pop(next(iter(self.cache)))
        self.cache[key] = value
    
    def clear(self) -> None:
        """Clear the cache."""
        self.cache.clear()
        
    def get_stats(self) -> Dict[str, int]:
        """Get cache statistics."""
        return {
            "hits": self.hits,
            "misses": self.misses,
            "size": len(self.cache),
            "max_size": self.max_size,
            "hit_ratio": self.hits / (self.hits + self.misses) if (self.hits + self.misses) > 0 else 0
        }

class DocumentIndex:
    """Efficient index for quickly finding text in documents."""
    
    def __init__(self):
        """Initialize document index."""
        self.word_index = {}  # Maps words to paragraph/run positions
        self.phrase_cache = {}  # Cache for phrase lookups
        self.doc = None
    
    def index_document(self, doc: Document) -> None:
        """
        Build an index of all words in the document for faster searching.
        
        Args:
            doc: Document to index
        """
        self.doc = doc
        self.word_index = {}
        self.phrase_cache = {}
        
        # Index words in paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            words = re.findall(r'\b\w+\b', para.text.lower())
            for word in words:
                if word not in self.word_index:
                    self.word_index[word] = []
                self.word_index[word].append(('para', para_idx))
        
        # Index words in tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        words = re.findall(r'\b\w+\b', para.text.lower())
                        for word in words:
                            if word not in self.word_index:
                                self.word_index[word] = []
                            self.word_index[word].append(('table', table_idx, row_idx, cell_idx, para_idx))
    
    def find_text(self, text: str) -> List[Tuple]:
        """
        Find all occurrences of text in the document.
        
        Args:
            text: Text to find
            
        Returns:
            List of locations (paragraph or table indices)
        """
        # Check cache first
        if text in self.phrase_cache:
            return self.phrase_cache[text]
        
        # For single words, use the index directly
        text_