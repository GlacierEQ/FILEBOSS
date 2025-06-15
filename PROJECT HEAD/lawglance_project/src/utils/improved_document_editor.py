"""
Improved document editing module with better error handling, typing, and documentation.
This addresses issues with the existing DocumentEditor implementation.
"""
import os
import re
import logging
from typing import Dict, Any, Optional, List, Union, Callable, Tuple
from enum import Enum
from pathlib import Path

# Import Word document handling libraries
import docx
from docx import Document
from docx.shared import Pt, RGBColor

# Constants for editor configuration
LOGGER_NAME = "lawglance.document_editor"
DEFAULT_BACKUP_SUFFIX = ".bak"


class EditOperation(Enum):
    """Enumeration of supported document edit operations."""
    REPLACE = "replace"
    INSERT = "insert"
    DELETE = "delete"
    FORMAT = "format"
    ADD_SECTION = "add_section"
    MOVE = "move"


class FormattingStyle(Enum):
    """Enumeration of supported formatting styles."""
    BOLD = "bold"
    ITALIC = "italic"
    UNDERLINE = "underline"
    HEADING1 = "heading1"
    HEADING2 = "heading2"
    HEADING3 = "heading3"
    RED = "red"
    BLUE = "blue"
    GREEN = "green"


class DocumentEditorError(Exception):
    """Base exception for document editor errors."""
    pass


class UnsupportedFormatError(DocumentEditorError):
    """Exception raised when a document format is not supported."""
    pass


class ParseInstructionError(DocumentEditorError):
    """Exception raised when instructions cannot be parsed."""
    pass


class EditOperationError(DocumentEditorError):
    """Exception raised when an edit operation fails."""
    pass


class ImprovedDocumentEditor:
    """
    Enhanced document editor that provides natural language editing capabilities
    with improved error handling, typing, and documentation.
    """

    def __init__(self) -> None:
        """Initialize the improved document editor."""
        self.logger = self._initialize_logger()
        self.com_available = self._check_com_availability()
        
        # Map edit operations to their handler methods
        self.edit_operations: Dict[EditOperation, Callable] = {
            EditOperation.REPLACE: self._replace_text,
            EditOperation.INSERT: self._insert_text,
            EditOperation.DELETE: self._delete_text,
            EditOperation.FORMAT: self._format_text,
            EditOperation.ADD_SECTION: self._add_section,
            EditOperation.MOVE: self._move_text
        }
        
        # Map formatting styles to their handler methods
        self.formatting_styles: Dict[FormattingStyle, Callable] = {
            FormattingStyle.BOLD: self._apply_bold,
            FormattingStyle.ITALIC: self._apply_italic,
            FormattingStyle.UNDERLINE: self._apply_underline,
            FormattingStyle.HEADING1: lambda elem: self._apply_heading_style(elem, 1),
            FormattingStyle.HEADING2: lambda elem: self._apply_heading_style(elem, 2),
            FormattingStyle.HEADING3: lambda elem: self._apply_heading_style(elem, 3),
            FormattingStyle.RED: self._apply_red_color,
            FormattingStyle.BLUE: self._apply_blue_color,
            FormattingStyle.GREEN: self._apply_green_color,
        }

    def _initialize_logger(self) -> logging.Logger:
        """
        Initialize and configure the logger.
        
        Returns:
            Configured logger instance
        """
        logger = logging.getLogger(LOGGER_NAME)
        
        # Only add handlers if none exist
        if not logger.handlers:
            handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            handler.setFormatter(formatter)
            logger.addHandler(handler)
            logger.setLevel(logging.INFO)
            
        return logger

    def _check_com_availability(self) -> bool:
        """
        Check if Windows COM automation is available for advanced document handling.
        
        Returns:
            True if COM automation is available, False otherwise
        """
        try:
            import win32com.client
            import pythoncom
            return True
        except ImportError:
            self.logger.info("Windows COM automation not available. Some features will be limited.")
            return False
        except Exception as e:
            self.logger.warning(f"Error checking COM availability: {str(e)}")
            return False

    def edit_document(self, file_path: str, instructions: str) -> Tuple[bool, str]:
        """
        Edit a document based on natural language instructions.
        
        Args:
            file_path: Path to the document to edit
            instructions: Natural language instructions for editing
            
        Returns:
            Tuple of (success, message)
            
        Raises:
            UnsupportedFormatError: If the document format is not supported
            ParseInstructionError: If the instructions cannot be parsed
            EditOperationError: If the edit operation fails
        """
        # Validate file format
        if not self._is_supported_format(file_path):
            raise UnsupportedFormatError(f"Unsupported document format: {Path(file_path).suffix}")
            
        try:
            # Create backup
            backup_path = self._create_backup(file_path)
            
            # Parse instructions
            operation, params = self._parse_instructions(instructions)
            
            # Load document
            doc = self._load_document(file_path)
            
            # Execute edit operation
            if operation not in self.edit_operations:
                raise ParseInstructionError(f"Unsupported operation: {operation}")
                
            handler = self.edit_operations[operation]
            success, message = handler(doc, **params)
            
            # Save document if edit was successful
            if success:
                doc.save(file_path)
                
            return success, message
            
        except ParseInstructionError as e:
            self.logger.error(f"Failed to parse instructions: {str(e)}")
            raise
        except docx.opc.exceptions.PackageNotFoundError:
            self.logger.error(f"Document not found or not accessible: {file_path}")
            raise EditOperationError(f"Document not found or not accessible: {file_path}")
        except Exception as e:
            self.logger.error(f"Error editing document: {str(e)}")
            raise EditOperationError(f"Error editing document: {str(e)}")

    def _is_supported_format(self, file_path: str) -> bool:
        """
        Check if the document format is supported.
        
        Args:
            file_path: Path to the document
            
        Returns:
            True if format is supported, False otherwise
        """
        supported_formats = ['.docx']
        if self.com_available:
            supported_formats.append('.doc')
            
        file_ext = Path(file_path).suffix.lower()
        return file_ext in supported_formats

    def _create_backup(self, file_path: str) -> str:
        """
        Create a backup of the document before editing.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Path to the backup file
        """
        backup_path = file_path + DEFAULT_BACKUP_SUFFIX
        try:
            import shutil
            shutil.copy2(file_path, backup_path)
            self.logger.info(f"Created backup at {backup_path}")
            return backup_path
        except Exception as e:
            self.logger.warning(f"Failed to create backup: {str(e)}")
            return ""

    def _load_document(self, file_path: str) -> Document:
        """
        Load a document from file.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Loaded document object
            
        Raises:
            EditOperationError: If the document cannot be loaded
        """
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.docx':
                return Document(file_path)
                
            elif file_ext == '.doc' and self.com_available:
                # Convert .doc to .docx first
                import win32com.client
                import pythoncom
                import tempfile
                
                pythoncom.CoInitialize()
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp:
                        temp_path = temp.name
                    
                    # Open and save as docx
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    doc.SaveAs2(temp_path, FileFormat=16)  # 16 = docx
                    doc.Close()
                    word.Quit()
                    
                    # Load the docx file
                    result = Document(temp_path)
                    
                    # Clean up
                    os.unlink(temp_path)
                    return result
                except Exception as e:
                    raise EditOperationError(f"Failed to convert .doc file: {str(e)}")
                finally:
                    pythoncom.CoUninitialize()
            else:
                raise UnsupportedFormatError(f"Unsupported document format: {file_ext}")
                
        except Exception as e:
            raise EditOperationError(f"Failed to load document: {str(e)}")

    def _parse_instructions(self, instructions: str) -> Tuple[EditOperation, Dict[str, Any]]:
        """
        Parse natural language instructions into an operation and parameters.
        
        Args:
            instructions: Natural language instructions
            
        Returns:
            Tuple of (operation, parameters)
            
        Raises:
            ParseInstructionError: If instructions cannot be parsed
        """
        if not instructions:
            raise ParseInstructionError("Instructions cannot be empty")
            
        instructions = instructions.lower()
        
        try:
            # Replace operation
            replace_match = re.search(r"replace\s+['\"](.+?)['\"]\s+with\s+['\"](.+?)['\"]", instructions)
            if replace_match:
                return EditOperation.REPLACE, {
                    "old_text": replace_match.group(1),
                    "new_text": replace_match.group(2)
                }
                
            # Insert operation
            insert_match = re.search(r"insert\s+['\"](.+?)['\"]\s+(?:at|after|before)\s+['\"](.+?)['\"]", instructions)
            if insert_match:
                return EditOperation.INSERT, {
                    "text": insert_match.group(1),
                    "position": insert_match.group(2)
                }
                
            # Insert at beginning/end
            begin_match = re.search(r"insert\s+['\"](.+?)['\"](?:\s+at)?\s+(?:the\s+)?beginning", instructions)
            if begin_match:
                return EditOperation.INSERT, {
                    "text": begin_match.group(1),
                    "position": "start"
                }
                
            end_match = re.search(r"insert\s+['\"](.+?)['\"](?:\s+at)?\s+(?:the\s+)?end", instructions)
            if end_match:
                return EditOperation.INSERT, {
                    "text": end_match.group(1),
                    "position": "end"
                }
                
            # Delete operation
            delete_match = re.search(r"(?:delete|remove)\s+['\"](.+?)['\"]", instructions)
            if delete_match:
                return EditOperation.DELETE, {
                    "text": delete_match.group(1)
                }
                
            # Format operation
            format_match = re.search(r"(?:format|make|set)\s+['\"](.+?)['\"]\s+(?:as|to)\s+(\w+)", instructions)
            if format_match:
                try:
                    style = FormattingStyle(format_match.group(2))
                    return EditOperation.FORMAT, {
                        "text": format_match.group(1),
                        "style": style
                    }
                except ValueError:
                    # If not a valid enum value, use the string
                    return EditOperation.FORMAT, {
                        "text": format_match.group(1),
                        "style": format_match.group(2)
                    }
                
            # Add section operation
            section_match = re.search(r"add\s+(?:a\s+)?(?:new\s+)?section\s+['\"](.+?)['\"]\s+with\s+content\s+['\"](.+?)['\"]", instructions)
            if section_match:
                return EditOperation.ADD_SECTION, {
                    "title": section_match.group(1),
                    "content": section_match.group(2)
                }
                
            # Move operation
            move_match = re.search(r"move\s+['\"](.+?)['\"]\s+(before|after)\s+['\"](.+?)['\"]", instructions)
            if move_match:
                return EditOperation.MOVE, {
                    "text": move_match.group(1),
                    "position": move_match.group(2),
                    "target": move_match.group(3)
                }
                
            raise ParseInstructionError(f"Could not understand instructions: {instructions}")
            
        except ParseInstructionError:
            raise
        except Exception as e:
            raise ParseInstructionError(f"Error parsing instructions: {str(e)}")

    def _replace_text(self, doc: Document, old_text: str, new_text: str) -> Tuple[bool, str]:
        """
        Replace text in a document.
        
        Args:
            doc: Document object
            old_text: Text to replace
            new_text: Replacement text
            
        Returns:
            Tuple of (success, message)
        """
        if not old_text:
            return False, "Text to replace cannot be empty"
            
        replacement_count = 0
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            if old_text in para.text:
                for run in para.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
                        replacement_count += 1
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            for run in para.runs:
                                if old_text in run.text:
                                    run.text = run.text.replace(old_text, new_text)
                                    replacement_count += 1
        
        if replacement_count > 0:
            return True, f"Replaced '{old_text}' with '{new_text}' {replacement_count} times"
        else:
            return False, f"Text '{old_text}' not found in document"

    def _insert_text(self, doc: Document, text: str, position: str = "end") -> Tuple[bool, str]:
        """
        Insert text at a specific position in document.
        
        Args:
            doc: Document object
            text: Text to insert
            position: Where to insert (start, end, or specific text)
            
        Returns:
            Tuple of (success, message)
        """
        # Implementation details would go here
        # This would handle special positions like "start" and "end"
        # as well as inserting after specific text
        pass
        
    def _delete_text(self, doc: Document, text: str) -> Tuple[bool, str]:
        """
        Delete text from document.
        
        Args:
            doc: Document object
            text: Text to delete
            
        Returns:
            Tuple of (success, message)
        """
        # Implementation details would go here
        pass
        
    def _format_text(self, doc: Document, text: str, style: Union[str, FormattingStyle]) -> Tuple[bool, str]:
        """
        Apply formatting to text.
        
        Args:
            doc: Document object
            text: Text to format
            style: Formatting style to apply
            
        Returns:
            Tuple of (success, message)
        """
        # Implementation details would go here
        pass
        
    def _add_section(self, doc: Document, title: str, content: str) -> Tuple[bool, str]:
        """
        Add a new section to document.
        
        Args:
            doc: Document object
            title: Section title
            content: Section content
            
        Returns:
            Tuple of (success, message)
        """
        # Implementation details would go here
        pass
        
    def _move_text(self, doc: Document, text: str, target: str, position: str) -> Tuple[bool, str]:
        """
        Move text to a new position.
        
        Args:
            doc: Document object
            text: Text to move
            target: Text to move relative to
            position: Before or after the target
            
        Returns:
            Tuple of (success, message)
        """
        # Implementation details would go here
        pass
        
    def _apply_bold(self, element: Any) -> None:
        """Apply bold formatting to a run or paragraph."""
        if hasattr(element, 'bold'):
            element.bold = True
        elif hasattr(element, 'runs') and element.runs:
            for run in element.runs:
                run.bold = True
                
    def _apply_italic(self, element: Any) -> None:
        """Apply italic formatting to a run or paragraph."""
        if hasattr(element, 'italic'):
            element.italic = True
        elif hasattr(element, 'runs') and element.runs:
            for run in element.runs:
                run.italic = True
                
    def _apply_underline(self, element: Any) -> None:
        """Apply underline formatting to a run or paragraph."""
        if hasattr(element, 'underline'):
            element.underline = True
        elif hasattr(element, 'runs') and element.runs:
            for run in element.runs:
                run.underline = True
                
    def _apply_heading_style(self, element: Any, level: int) -> None:
        """Apply heading style to a paragraph."""
        # Implementation details would go here
        pass
        
    def _apply_red_color(self, element: Any) -> None:
        """Apply red color to text."""
        # Implementation details would go here
        pass
        
    def _apply_blue_color(self, element: Any) -> None:
        """Apply blue color to text."""
        # Implementation details would go here
        pass
        
    def _apply_green_color(self, element: Any) -> None:
        """Apply green color to text."""
        # Implementation details would go here
        pass
