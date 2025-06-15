"""Advanced formatting utilities for document editing."""
import re
from typing import Dict, List, Tuple, Optional, Any, Union
from dataclasses import dataclass, field
import logging
from enum import Enum
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE

logger = logging.getLogger("lawglance.advanced_formatting")

# Constants for formatting options
DEFAULT_FONT_SIZE = 11
DEFAULT_FONT_NAME = "Times New Roman"
HEADING_SIZES = {
    1: 16,  # Heading 1
    2: 14,  # Heading 2
    3: 12,  # Heading 3
    4: 11,  # Heading 4
}


class TextAlignment(Enum):
    """Enum for text alignment options."""
    LEFT = WD_ALIGN_PARAGRAPH.LEFT
    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    RIGHT = WD_ALIGN_PARAGRAPH.RIGHT
    JUSTIFY = WD_ALIGN_PARAGRAPH.JUSTIFY


class UnderlineStyle(Enum):
    """Enum for underline style options."""
    SINGLE = WD_UNDERLINE.SINGLE
    DOUBLE = WD_UNDERLINE.DOUBLE
    DOTTED = WD_UNDERLINE.DOTTED
    DASH = WD_UNDERLINE.DASH
    NONE = WD_UNDERLINE.NONE


@dataclass
class ColorScheme:
    """Class representing a color scheme."""
    primary: RGBColor
    secondary: RGBColor
    accent: RGBColor
    
    @classmethod
    def legal_blue(cls) -> 'ColorScheme':
        """Legal blue color scheme."""
        return cls(
            primary=RGBColor(0, 51, 102),   # Dark blue
            secondary=RGBColor(100, 149, 237), # Cornflower blue
            accent=RGBColor(192, 192, 192)  # Silver
        )
    
    @classmethod
    def legal_maroon(cls) -> 'ColorScheme':
        """Legal maroon color scheme."""
        return cls(
            primary=RGBColor(128, 0, 0),     # Maroon
            secondary=RGBColor(139, 0, 0),   # Dark red
            accent=RGBColor(220, 220, 220)   # Light gray
        )


@dataclass
class AdvancedFormatSpec:
    """Advanced formatting specification."""
    # Font properties
    font_name: Optional[str] = None
    font_size: Optional[int] = None
    font_color: Optional[RGBColor] = None
    
    # Emphasis properties
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[UnderlineStyle] = None
    
    # Paragraph properties
    alignment: Optional[TextAlignment] = None
    line_spacing: Optional[float] = None
    space_before: Optional[int] = None
    space_after: Optional[int] = None
    
    # Style properties
    style_name: Optional[str] = None
    heading_level: Optional[int] = None
    
    def apply_to_run(self, run: docx.text.run.Run) -> None:
        """Apply formatting to a run."""
        if self.font_name:
            run.font.name = self.font_name
        if self.font_size:
            run.font.size = Pt(self.font_size)
        if self.font_color:
            run.font.color.rgb = self.font_color
        if self.bold is not None:
            run.bold = self.bold
        if self.italic is not None:
            run.italic = self.italic
        if self.underline:
            run.underline = self.underline.value
    
    def apply_to_paragraph(self, paragraph: docx.text.paragraph.Paragraph) -> None:
        """Apply formatting to a paragraph."""
        if self.alignment:
            paragraph.alignment = self.alignment.value
        if self.line_spacing:
            paragraph.paragraph_format.line_spacing = self.line_spacing
        if self.space_before:
            paragraph.paragraph_format.space_before = Pt(self.space_before)
        if self.space_after:
            paragraph.paragraph_format.space_after = Pt(self.space_after)
        if self.style_name:
            try:
                paragraph.style = self.style_name
            except Exception as e:
                logger.warning(f"Could not apply style '{self.style_name}': {e}")
        if self.heading_level and 0 < self.heading_level <= 9:
            try:
                paragraph.style = f'Heading {self.heading_level}'
                # Set font size if not specified but heading level is
                if not self.font_size and self.heading_level in HEADING_SIZES:
                    for run in paragraph.runs:
                        run.font.size = Pt(HEADING_SIZES[self.heading_level])
            except Exception as e:
                logger.warning(f"Could not apply heading style: {e}")


class AdvancedFormatter:
    """Advanced document formatting utility."""
    
    def __init__(self):
        """Initialize the advanced formatter."""
        self.logger = logging.getLogger(f"{__name__}.{self.__class__.__name__}")
        
        # Pre-defined formats
        self.predefined_formats = {
            "legal_heading": AdvancedFormatSpec(
                font_name="Arial",
                font_size=14,
                bold=True,
                font_color=RGBColor(0, 51, 102),
                alignment=TextAlignment.LEFT,
                space_after=12
            ),
            "section_title": AdvancedFormatSpec(
                font_name="Times New Roman",
                font_size=12,
                bold=True,
                font_color=RGBColor(0, 0, 0),
                underline=UnderlineStyle.SINGLE,
                space_before=12,
                space_after=6
            ),
            "citation": AdvancedFormatSpec(
                font_name="Times New Roman",
                font_size=11,
                italic=True
            ),
            "body_text": AdvancedFormatSpec(
                font_name="Times New Roman",
                font_size=11,
                alignment=TextAlignment.JUSTIFY,
                line_spacing=1.15
            ),
            "footer": AdvancedFormatSpec(
                font_name="Arial",
                font_size=9,
                alignment=TextAlignment.CENTER
            )
        }
    
    def format_text(self, doc: Document, text: str, format_name: str) -> int:
        """
        Apply a predefined format to text in a document.
        
        Args:
            doc: Document to edit
            text: Text to format
            format_name: Name of predefined format
            
        Returns:
            Number of formatting operations performed
        """
        if format_name not in self.predefined_formats:
            self.logger.error(f"Format '{format_name}' not found in predefined formats")
            return 0
            
        format_spec = self.predefined_formats[format_name]
        return self._apply_format(doc, text, format_spec)
    
    def create_custom_format(self, **kwargs) -> AdvancedFormatSpec:
        """
        Create a custom format specification.
        
        Args:
            **kwargs: Format settings
            
        Returns:
            AdvancedFormatSpec with specified settings
        """
        return AdvancedFormatSpec(**kwargs)
    
    def apply_custom_format(self, doc: Document, text: str, format_spec: AdvancedFormatSpec) -> int:
        """
        Apply a custom format to text in a document.
        
        Args:
            doc: Document to edit
            text: Text to format
            format_spec: Format specification
            
        Returns:
            Number of formatting operations performed
        """
        return self._apply_format(doc, text, format_spec)
    
    def apply_color_scheme(self, doc: Document, scheme: ColorScheme) -> None:
        """
        Apply a color scheme to a document.
        
        Args:
            doc: Document to edit
            scheme: Color scheme to apply
        """
        # Apply to heading styles
        for i in range(1, 10):  # Heading 1-9
            try:
                heading_style = doc.styles[f'Heading {i}']
                if i == 1:
                    heading_style.font.color.rgb = scheme.primary
                elif i == 2:
                    heading_style.font.color.rgb = scheme.secondary
                else:
                    heading_style.font.color.rgb = scheme.secondary
            except Exception as e:
                self.logger.debug(f"Could not update style 'Heading {i}': {e}")
    
    def _apply_format(self, doc: Document, text: str, format_spec: AdvancedFormatSpec) -> int:
        """
        Apply format to text in document.
        
        Args:
            doc: Document to edit
            text: Text to format
            format_spec: Format to apply
            
        Returns:
            Number of formatting operations performed
        """
        format_count = 0
        
        # Format paragraphs
        for para in doc.paragraphs:
            if text in para.text:
                exact_match = text == para.text.strip()
                
                # Apply paragraph formatting for exact matches or if certain paragraph settings exist
                if exact_match or any([
                    format_spec.alignment, 
                    format_spec.line_spacing, 
                    format_spec.space_before, 
                    format_spec.space_after,
                    format_spec.style_name,
                    format_spec.heading_level
                ]):
                    format_spec.apply_to_paragraph(para)
                    format_count += 1
                
                # Apply run formatting
                for run in para.runs:
                    if text in run.text:
                        if exact_match:
                            # Format entire run
                            format_spec.apply_to_run(run)
                            format_count += 1
                        else:
                            # For partial matches, need to split the run and format only the matching text
                            format_count += self._format_partial_match(para, run, text, format_spec)
        
        # Format tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if text in para.text:
                            exact_match = text == para.text.strip()
                            
                            # Apply paragraph formatting
                            if exact_match or any([
                                format_spec.alignment, 
                                format_spec.line_spacing, 
                                format_spec.space_before, 
                                format_spec.space_after,
                                format_spec.style_name,
                                format_spec.heading_level
                            ]):
                                format_spec.apply_to_paragraph(para)
                                format_count += 1
                            
                            # Apply run formatting
                            for run in para.runs:
                                if text in run.text:
                                    if exact_match:
                                        format_spec.apply_to_run(run)
                                        format_count += 1
                                    else:
                                        format_count += self._format_partial_match(para, run, text, format_spec)
                                        
        return format_count
    
    def _format_partial_match(self, para, run, text, format_spec):
        """
        Format only the matching part of a run.
        
        This is a complex operation as we need to split the run, which requires rebuilding the paragraph.
        
        Args:
            para: Paragraph containing the run
            run: Run containing the text
            text: Text to format
            format_spec: Format to apply
            
        Returns:
            Number of formatting operations performed (0 or 1)
        """
        # Not implementing this complex logic in this example, as it requires detailed
        # handling of run splitting which is beyond the scope of this answer
        self.logger.warning("Partial match formatting not implemented")
        return 0
