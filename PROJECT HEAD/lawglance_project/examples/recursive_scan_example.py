"""
Example demonstrating the recursive document scanning functionality.
This shows how to deeply analyze document structure, metadata, and nested elements.
"""
import os
import sys
from pathlib import Path

# Add project root to path
sys.path.append(str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt
from src.utils.document_editor import DocumentEditor

def create_complex_document(path):
    """Create a document with complex structure to test recursive scanning."""
    doc = Document()
    
    # Add title
    doc.add_heading("Complex Document with Nested Elements", 0)
    
    # Add metadata
    doc.core_properties.author = "Test Author"
    # Intentionally leave title blank to trigger metadata issue
    doc.core_properties.comments = "CONFIDENTIAL: This contains sensitive information"
    
    # Add some normal paragraphs
    doc.add_paragraph("This is a standard paragraph with normal formatting.")
    
    # Add paragraph with formatting issues
    p = doc.add_paragraph("This paragraph has ")
    p.add_run("inconsistent").bold = True
    p.add_run(" formatting ")
    run = p.add_run("with")
    run.font.name = "Arial"
    p.add_run(" multiple ")
    run = p.add_run("fonts")
    run.font.name = "Courier New"
    p.add_run(" used.")
    
    # Add a heading with skipped level (skipping level 1)
    doc.add_heading("This is a heading level 2", 2)
    
    # Add paragraph with double  spaces  like  this
    doc.add_paragraph("This  paragraph  has  multiple  double  spaces.")
    
    # Add an empty paragraph
    doc.add_paragraph("")
    
    # Add a simple table
    table = doc.add_table(rows=3, cols=3)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f"Cell {i+1},{j+1}"
    
    # Add a nested table (table inside a table cell)
    table = doc.add_table(rows=2, cols=2)
    # Put a nested table in the first cell
    nested_cell = table.cell(0, 0)
    nested_cell.text = ""  # Clear default content
    nested_table = nested_cell.add_table(rows=2, cols=2)
    nested_table.cell(0, 0).text = "Nested cell 1,1"
    nested_table.cell(0, 1).text = "Nested cell 1,2"
    nested_table.cell(1, 0).text = "Nested cell 2,1"
    nested_table.cell(1, 1).text = "Nested cell 2,2"
    
    # Fill other cells
    table.cell(0, 1).text = "Normal cell"
    table.cell(1, 0).text = "Normal cell"
    table.cell(1, 1).text = "Normal cell"
    
    # Add a jagged table (inconsistent columns)
    table = doc.add_table(rows=3, cols=0)
    row_cells = table.add_row().cells
    for _ in range(2):
        row_cells.add_paragraph("Row 1 cell")
        
    row_cells = table.add_row().cells
    for _ in range(3):  # Different number of cells
        row_cells.add_paragraph("Row 2 cell")
    
    row_cells = table.add_row().cells
    for _ in range(4):  # Different number of cells
        row_cells.add_paragraph("Row 3 cell")
    
    # Add an extremely long paragraph
    long_text = "This is a very long paragraph. " * 100
    doc.add_paragraph(long_text)
    
    # Add text with unusual font sizes
    p = doc.add_paragraph()
    run = p.add_run("This text is very small. ")
    run.font.size = Pt(6)
    run = p.add_run("This text is normal. ")
    run.font.size = Pt(11)
    run = p.add_run("This text is very large.")
    run.font.size = Pt(36)
    
    # Save the document
    doc.save(path)
    print(f"Created complex document at {path}")

def main():
    """Run the recursive scan example."""
    print("Recursive Document Scan Example")
    print("=" * 50)
    
    # Create path for our test document
    sample_path = Path(__file__).parent.parent / "data" / "complex_document.docx"
    sample_path.parent.mkdir(exist_ok=True)
    
    # Create a document with various issues to detect
    create_complex_document(sample_path)
    
    # Initialize document editor
    editor = DocumentEditor()
    doc = Document(sample_path)
    
    # Perform a regular scan first
    print("\n1. Basic Document Scan:")
    print("-" * 30)
    result = editor._scan_for_repairs(doc)
    print(result)
    
    # Now perform a recursive deep scan
    print("\n2. Recursive Deep Scan:")
    print("-" * 30)
    result = editor._recursive_scan(doc)
    print(result)
    
    # Try with auto-fix enabled
    print("\n3. Recursive Scan with Auto-Fix:")
    print("-" * 30)
    result = editor._recursive_scan(doc, auto_fix=True)
    print(result)
    
    # Save the potentially fixed document
    fixed_path = Path(__file__).parent.parent / "data" / "complex_document_fixed.docx"
    doc.save(fixed_path)
    print(f"\nSaved document after attempted fixes to: {fixed_path}")
    
if __name__ == "__main__":
    main()
