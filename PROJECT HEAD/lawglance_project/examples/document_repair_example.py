"""
Example of using the document editor's repair functionality.
This demonstrates how to scan documents for issues and fix them.
"""

import sys
from pathlib import Path

from pathlib import Path

# Add project root to path
sys.path.append(str(Path(__file__).parent.parent))

from docx import Document
from src.utils.document_editor import DocumentEditor
from src.utils.enhanced_document_editor import EnhancedDocumentEditor

def create_sample_document(path):
    """Create a sample document with some issues to repair."""
    doc = Document()
    
    # Add some content with various issues
    doc.add_heading("Sample  Document with  Issues", 0)  # Double spaces
    
    # Empty paragraph
    doc.add_paragraph("")
    
    # Inconsistent formatting
    para = doc.add_paragraph("This paragraph has ")
    para.add_run("different").bold = True
    para.add_run(" fonts ")
    run = para.add_run("mixed")
    run.font.name = "Arial"
    para.add_run(" together.")
    
    # Numbering issues
    doc.add_paragraph("1. First item")
    doc.add_paragraph("3. Third item")  # Skips 2
    
    # Some duplicate content
    doc.add_paragraph("This is duplicate content that should be detected.")
    doc.add_paragraph("Some other text here.")
    doc.add_paragraph("This is duplicate content that should be detected.")
    
    doc.save(path)
    print(f"Created sample document at {path}")

def main():
    """Run an example of document repair functionality."""
    print("Document Repair Example")
    print("=" * 40)
    
    # Create a path for our sample document
    sample_path = Path(__file__).parent.parent / "data" / "document_with_issues.docx"
    sample_path.parent.mkdir(exist_ok=True)
    
    # Create a sample document with issues
    create_sample_document(sample_path)
    
    print("\n1. Using basic DocumentEditor scan_for_repairs with detailed output:")

    basic_editor = DocumentEditor()
    doc1 = Document(sample_path)
    scan_result = basic_editor.scan_for_repairs(doc1)  # Updated to use public method












    print(scan_result)
    
    print("\n2. Using EnhancedDocumentEditor with full repair functionality and detailed output:")

    enhanced_editor = EnhancedDocumentEditor()
    doc2 = Document(sample_path)
    print("\nScanning document...")
    scan_result = enhanced_editor.scan_for_repairs(doc2)  # Updated to use public method


    
    # Scan for issues
    print("\nScanning document...")
    scan_result = enhanced_editor.scan_for_repairs(doc2)  # Updated to use public method

    print("Scan Results:", scan_result)

    
    # Repair issues
    print("\nRepairing document...")
    repair_result = enhanced_editor.repair_document(doc2)  # Updated to use public method












    print("Repair Results:", repair_result)

    
    # Save the repaired document
    repaired_path = Path(__file__).parent.parent / "data" / "document_repaired.docx"
    doc2.save(repaired_path)
    print(f"\nSuccessfully saved repaired document to {repaired_path}")

    
if __name__ == "__main__":
    main()
