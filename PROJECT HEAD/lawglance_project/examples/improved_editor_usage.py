"""
Example demonstrating the improved document editor with proper error handling and typing.
This example shows how to properly use the ImprovedDocumentEditor class.
"""
import os
import sys
import logging
from pathlib import Path
from typing import Optional, Tuple

# Constants
PROJECT_ROOT = Path(__file__).parent.parent
SAMPLE_DOC_PATH = PROJECT_ROOT / "data" / "sample_document.docx"
LOG_NAME = "lawglance.examples.editor_usage"

# Add project root to path if needed
sys.path.append(str(PROJECT_ROOT))

# Import the improved document editor
from src.utils.improved_document_editor import (
    ImprovedDocumentEditor, 
    UnsupportedFormatError,
    ParseInstructionError,
    EditOperationError
)
from src.utils.logging_setup import get_logger

# Initialize logger using the centralized logging system
logger = get_logger(LOG_NAME)


def create_sample_document() -> str:
    """
    Create a sample document if it doesn't exist.
    
    Returns:
        Path to the sample document
    """
    # Create directory if it doesn't exist
    SAMPLE_DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    
    # Only create if it doesn't exist
    if not SAMPLE_DOC_PATH.exists():
        try:
            from docx import Document
            
            # Create a new document
            doc = Document()
            doc.add_heading("Sample Document", 0)
            doc.add_paragraph("This is a sample paragraph to edit.")
            doc.add_paragraph("This text will be edited with the ImprovedDocumentEditor.")
            doc.add_paragraph("The editor has improved error handling and type annotations.")
            
            # Save the document
            doc.save(SAMPLE_DOC_PATH)
            logger.info(f"Created sample document at {SAMPLE_DOC_PATH}")
        except Exception as e:
            logger.error(f"Failed to create sample document: {str(e)}")
            return ""
    
    return str(SAMPLE_DOC_PATH)


def process_edit_instruction(
    file_path: str, 
    instruction: str
) -> Tuple[bool, str]:
    """
    Process an edit instruction on a document with comprehensive error handling.
    
    Args:
        file_path: Path to document to edit
        instruction: Natural language instruction for editing
        
    Returns:
        Tuple of (success, message)
    """
    editor = ImprovedDocumentEditor()
    
    try:
        # Validate file exists
        if not Path(file_path).exists():
            return False, f"File not found: {file_path}"
            
        # Process the edit instruction
        logger.info(f"Processing edit instruction: '{instruction}'")
        success, message = editor.edit_document(file_path, instruction)
        
        if success:
            logger.info(f"Successfully edited document: {message}")
        else:
            logger.warning(f"Edit unsuccessful: {message}")
            
        return success, message
        
    except UnsupportedFormatError as e:
        logger.error(f"Unsupported file format: {str(e)}")
        return False, f"Unsupported file format: {str(e)}"
        
    except ParseInstructionError as e:
        logger.error(f"Failed to parse instruction: {str(e)}")
        return False, f"Failed to understand instruction: {str(e)}"
        
    except EditOperationError as e:
        logger.error(f"Edit operation failed: {str(e)}")
        return False, f"Edit operation failed: {str(e)}"
        
    except PermissionError as e:
        logger.error(f"Permission error: {str(e)}")
        return False, f"Cannot access file. Ensure it's not open in another program."
        
    except Exception as e:
        logger.exception(f"Unexpected error: {str(e)}")
        return False, f"An unexpected error occurred: {str(e)}"


def demonstration() -> None:
    """Run a demonstration of the document editor with various instructions."""
    print("ImprovedDocumentEditor Demonstration")
    print("=" * 40)
    
    # Create or get the sample document
    sample_path = create_sample_document()
    if not sample_path:
        print("Failed to create or locate sample document.")
        return
    
    print(f"Using sample document at: {sample_path}")
    print()
    
    # List of edit instructions to demonstrate
    instructions = [
        "replace 'sample paragraph' with 'demonstration paragraph'",
        "format 'Sample Document' as heading1",
        "insert 'This text was inserted by the editor.' at the end",
        "delete 'to edit'",
    ]
    
    # Process each instruction
    for i, instruction in enumerate(instructions, 1):
        print(f"Instruction {i}: {instruction}")
        success, message = process_edit_instruction(sample_path, instruction)
        
        if success:
            print(f"✓ Success: {message}")
        else:
            print(f"✗ Failed: {message}")
        print()
    
    print("Demonstration complete.")
    print(f"You can open the document at {sample_path} to see the results.")


if __name__ == "__main__":
    demonstration()
