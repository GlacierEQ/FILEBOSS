"""
Simple example showing how to scan documents for issues using both basic and recursive methods.
"""
import os
import sys
from pathlib import Path
from docx import Document

# Add the project root to path for imports
project_root = Path(__file__).parent.parent
sys.path.append(str(project_root))

# Import the document editor
from src.utils.document_editor import DocumentEditor

def scan_document(doc_path):
    """
    Scan a document for issues using both regular and recursive methods.
    
    Args:
        doc_path: Path to the document to scan
    """
    # Make sure the document exists
    if not os.path.exists(doc_path):
        print(f"Error: Document not found at {doc_path}")
        return
        
    print(f"Scanning document: {doc_path}")
    print("=" * 50)
    
    # Load the document
    doc = Document(doc_path)
    
    # Create document editor
    editor = DocumentEditor()
    
    # Method 1: Basic scan for common issues
    print("\n1. Basic Document Scan:")
    print("-" * 30)
    basic_result = editor._scan_for_repairs(doc)
    print(basic_result)
    
    # Method 2: Thorough recursive scan
    print("\n2. Deep Recursive Scan:")
    print("-" * 30)
    recursive_result = editor._recursive_scan(doc)
    print(recursive_result)
    
    # Method 3: Scan with auto-fix
    print("\n3. Scan with Auto-Fix:")
    print("-" * 30)
    autofix_result = editor._recursive_scan(doc, auto_fix=True)
    print(autofix_result)
    
    # Save the fixed document
    fixed_path = os.path.splitext(doc_path)[0] + "_fixed.docx"
    doc.save(fixed_path)
    print(f"\nSaved fixed document to: {fixed_path}")

def main():
    """Main function to demonstrate document scanning."""
    # Check if a file path was provided as an argument
    if len(sys.argv) > 1:
        document_path = sys.argv[1]
    else:
        # Default to a sample document in the data directory
        document_path = project_root / "data" / "sample_legal_text.docx"
        
        # If the sample document doesn't exist, try to create one
        if not document_path.exists():
            try:
                # Create a simple document with some issues
                doc = Document()
                doc.add_heading("Sample Document With Issues", 0)
                doc.add_paragraph("This  paragraph  has  double  spaces.")
                doc.add_paragraph("")  # Empty paragraph
                
                # Add inconsistent formatting
                p = doc.add_paragraph("This paragraph has ")
                p.add_run("different").bold = True
                p.add_run(" font ")
                r = p.add_run("formatting")
                r.font.name = "Arial"
                
                # Create the directory if it doesn't exist
                document_path.parent.mkdir(exist_ok=True)
                
                # Save the document
                doc.save(document_path)
                print(f"Created sample document at {document_path}")
            except Exception as e:
                print(f"Error creating sample document: {str(e)}")
                return
    
    # Scan the document
    scan_document(document_path)

if __name__ == "__main__":
    main()
