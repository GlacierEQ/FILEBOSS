"""Test fixtures and configuration for pytest."""
import os
import tempfile
import shutil
from typing import Dict, Any, Tuple, Generator, List
import pytest
from docx import Document

# Constants for testing
TEST_CONSTANTS = {
    "SAMPLE_TEXT": "This is sample text for testing.",
    "SAMPLE_LEGAL_TEXT": (
        "CONTRACT AGREEMENT\n\n"
        "1. FIRST SECTION\nThis is the first section.\n\n"
        "2. SECOND SECTION\nThis is the second section.\n\n"
        "3. TERMINATION\nThis agreement may be terminated."
    ),
    "SIMILARITY_THRESHOLD": 0.8,
    "DEFAULT_TEST_TIMEOUT": 5  # seconds
}

@pytest.fixture(scope="function")
def temp_test_dir() -> Generator[str, None, None]:
    """
    Create a temporary directory for test files.
    
    Returns:
        Path to temporary directory, which will be cleaned up after test
    """
    # Create a temporary directory
    temp_dir = tempfile.mkdtemp()
    try:
        # Return the path to the temporary directory to the test function
        yield temp_dir
    finally:
        # Clean up the temporary directory
        shutil.rmtree(temp_dir)

@pytest.fixture(scope="function")
def sample_text_file(temp_test_dir: str) -> str:
    """
    Create a sample text file for testing.
    
    Args:
        temp_test_dir: Temporary directory path (from fixture)
        
    Returns:
        Path to the created text file
    """
    file_path = os.path.join(temp_test_dir, "sample.txt")
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(TEST_CONSTANTS["SAMPLE_LEGAL_TEXT"])
    return file_path

@pytest.fixture(scope="function")
def sample_docx_file(temp_test_dir: str) -> str:
    """
    Create a sample DOCX file for testing.
    
    Args:
        temp_test_dir: Temporary directory path (from fixture)
        
    Returns:
        Path to the created DOCX file
    """
    file_path = os.path.join(temp_test_dir, "sample.docx")
    doc = Document()
    
    # Add heading
    doc.add_heading("CONTRACT AGREEMENT", 0)
    
    # Add sections
    doc.add_heading("1. FIRST SECTION", level=1)
    doc.add_paragraph("This is the first section.")
    
    doc.add_heading("2. SECOND SECTION", level=1)
    doc.add_paragraph("This is the second section.")
    
    doc.add_heading("3. TERMINATION", level=1)
    doc.add_paragraph("This agreement may be terminated.")
    
    # Save the document
    doc.save(file_path)
    return file_path

@pytest.fixture(scope="function")
def similar_documents(temp_test_dir: str) -> Tuple[str, str]:
    """
    Create two similar documents with controlled differences.
    
    Args:
        temp_test_dir: Temporary directory path (from fixture)
        
    Returns:
        Tuple of paths to two similar documents
    """
    doc1_path = os.path.join(temp_test_dir, "doc1.txt")
    doc2_path = os.path.join(temp_test_dir, "doc2.txt")
    
    # Original document
    with open(doc1_path, "w", encoding="utf-8") as f:
        f.write("CONTRACT AGREEMENT\n\n")
        f.write("1. SECTION ONE\nThis is the first section.\n\n")
        f.write("2. SECTION TWO\nThis is the second section.\n\n")
        f.write("3. SECTION THREE\nThis is the third section.")
    
    # Similar document with controlled differences
    with open(doc2_path, "w", encoding="utf-8") as f:
        f.write("CONTRACT AGREEMENT\n\n")
        f.write("1. SECTION ONE\nThis is the modified first section.\n\n")
        f.write("2. SECTION TWO\nThis is the second section.\n\n")
        f.write("4. NEW SECTION\nThis is a new section.")
        
    return doc1_path, doc2_path

@pytest.fixture(scope="function")
def mock_document_info() -> Dict[str, Any]:
    """
    Create mock document information for testing.
    
    Returns:
        Dictionary with document metadata
    """
    return {
        "content": TEST_CONSTANTS["SAMPLE_LEGAL_TEXT"],
        "paragraphs": TEST_CONSTANTS["SAMPLE_LEGAL_TEXT"].split("\n\n"),
        "sections": ["CONTRACT AGREEMENT", "1. FIRST SECTION", "2. SECOND SECTION", "3. TERMINATION"],
        "word_count": len(TEST_CONSTANTS["SAMPLE_LEGAL_TEXT"].split())
    }

@pytest.fixture(scope="function")
def expected_comparison_result() -> Dict[str, Any]:
    """
    Create an expected comparison result for testing.
    
    Returns:
        Dictionary with expected comparison data
    """
    return {
        "success": True,
        "similarity_percentage": 85.5,  # Example value
        "difference_count": 3,  # Example value
        "additions": ["4. NEW SECTION", "This is a new section."],
        "deletions": ["3. SECTION THREE", "This is the third section."],
        "changes": [
            {
                "removed": "This is the first section.", 
                "added": "This is the modified first section."
            }
        ],
        "summary": "The documents are very similar with some differences."
    }
