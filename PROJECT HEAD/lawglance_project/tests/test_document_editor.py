"""Tests for the Document Editor module."""
import os
import sys
import unittest
import tempfile
import shutil
from docx import Document

# Add the src directory to the path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(project_root)

from src.utils.document_editor import DocumentEditor, TextEditor, FormattingEditor, InstructionParser

class TestInstructionParser(unittest.TestCase):
    """Test the instruction parser component."""
    
    def setUp(self):
        self.parser = InstructionParser()
    
    def test_parse_replace_instructions(self):
        """Test parsing replace instructions."""
        instructions = 'Replace "old text" with "new text"'
        operation, params = self.parser.parse_instructions(instructions)
        
        self.assertEqual(operation, "replace")
        self.assertEqual(params["old_text"], "old text")
        self.assertEqual(params["new_text"], "new text")
    
    def test_parse_insert_instructions(self):
        """Test parsing insert instructions."""
        instructions = 'Insert "new content" at end'
        operation, params = self.parser.parse_instructions(instructions)
        
        self.assertEqual(operation, "insert")
        self.assertEqual(params["text"], "new content")
        self.assertEqual(params["position"], "end")
    
    def test_parse_format_instructions(self):
        """Test parsing format instructions."""
        instructions = 'Make "important text" bold'
        operation, params = self.parser.parse_instructions(instructions)
        
        self.assertEqual(operation, "format")
        self.assertEqual(params["text"], "important text")
        self.assertEqual(params["style"], "bold")


class TestDocumentEditor(unittest.TestCase):
    """Test the document editor with actual files."""
    
    def setUp(self):
        """Create temporary test files."""
        self.editor = DocumentEditor()
        self.test_dir = tempfile.mkdtemp()
        
        # Create a test docx file
        self.test_docx = os.path.join(self.test_dir, "test_doc.docx")
        doc = Document()
        doc.add_paragraph("This is a test document. It contains sample text for testing.")
        doc.add_paragraph("Here is another paragraph with more test content.")
        doc.save(self.test_docx)
        
        # Create a test text file
        self.test_txt = os.path.join(self.test_dir, "test_file.txt")
        with open(self.test_txt, "w") as f:
            f.write("This is a test document.\nIt contains sample text for testing.\n")
            f.write("Here is another paragraph with more test content.\n")
    
    def tearDown(self):
        """Clean up test files."""
        shutil.rmtree(self.test_dir)
    
    def test_replace_text_docx(self):
        """Test replacing text in a docx document."""
        result = self.editor.edit_document(
            self.test_docx, 
            'Replace "test" with "sample"'
        )
        
        self.assertIn("Replaced", result)
        
        # Verify the change
        doc = Document(self.test_docx)
        text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("sample document", text)
        self.assertNotIn("test document", text)
    
    def test_insert_text_docx(self):
        """Test inserting text in a docx document."""
        result = self.editor.edit_document(
            self.test_docx, 
            'Insert "IMPORTANT: " at beginning'
        )
        
        self.assertIn("Inserted", result)
        
        # Verify the change
        doc = Document(self.test_docx)
        self.assertTrue(doc.paragraphs[0].text.startswith("IMPORTANT:"))
    
    def test_edit_text_file(self):
        """Test editing a text file."""
        result = self.editor.edit_document(
            self.test_txt, 
            'Replace "test" with "sample"'
        )
        
        self.assertIn("Successfully", result)
        
        # Verify the change
        with open(self.test_txt, "r") as f:
            content = f.read()
            self.assertIn("sample document", content)
            self.assertNotIn("test document", content)


if __name__ == "__main__":
    unittest.main()
